import streamlit as st
import streamlit.components.v1 as components
import os
import re
import json
import docx
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from dotenv import load_dotenv
from langdetect import detect, LangDetectException
from num2words import num2words
from openai import OpenAI
import openai

# BOMBA ATOMOWA NA INTERFEJS
st.set_page_config(page_title="Poliglota AI", layout="wide")

st.markdown("""
    <style>
        header {visibility: hidden !important;}
        footer {visibility: hidden !important;}
        #MainMenu {visibility: hidden !important;}
    </style>
""", unsafe_allow_html=True)

components.html("""
    <script>
        window.parent.document.documentElement.lang = 'pl';
    </script>
""", width=0, height=0)

# --- KROK 1: Bezpieczne inicjowanie środowiska ---
if os.path.exists("golden_key.env"):
    load_dotenv("golden_key.env")
    from openai import OpenAI
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    api_dostepne = True
else:
    client = None
    api_dostepne = False

st.title("Poliglota AI - Hybrydowe Studio Tłumaczeń i Akcentów")

def wyswietl_blad_ai(e, custom_msg=None):
    if custom_msg:
        st.error(custom_msg)
    else:
        st.error("🚨 Wystąpił nieoczekiwany błąd podczas przetwarzania. Szczegóły techniczne (do zgłoszenia) znajdziesz poniżej.")
    st.text_area("Treść błędu:", value=str(e), height=150)


if "file_content" not in st.session_state:
    st.session_state.file_content = ""
if "file_ext" not in st.session_state:
    st.session_state.file_ext = ""
if "target_lang_input" not in st.session_state:
    st.session_state.target_lang_input = ""
if "current_lang" not in st.session_state:
    st.session_state.current_lang = ""
if "ready_to_translate" not in st.session_state:
    st.session_state.ready_to_translate = False
if "oryginalna_nazwa" not in st.session_state:
    st.session_state.oryginalna_nazwa = "nieznany"
if "plik_katalog" not in st.session_state:
    st.session_state.plik_katalog = "."

# --- TWARDE REGUŁY FONETYCZNE (PYTHON) ---

def normalizuj_liczby(text):
    def zamien(match):
        try:
            return num2words(match.group(), lang='pl')
        except Exception:
            return match.group()
    return re.sub(r'\d+', zamien, text)

def sklej_pojedyncze_litery(text):
    return re.sub(r'(?i)\b([a-z])\s+', r'\1', text)

def usun_polskie_znaki(text):
    text = normalizuj_liczby(text)
    mapping = {
        'ą': 'on', 'ę': 'en', 'ł': 'l', 'ó': 'u',
        'ś': 's', 'ć': 'c', 'ń': 'n', 'ż': 'z', 'ź': 'z',
        'Ą': 'On', 'Ę': 'En', 'Ł': 'L', 'Ó': 'U',
        'Ś': 'S', 'Ć': 'C', 'Ń': 'N', 'Ż': 'Z', 'Ź': 'Z'
    }
    for k, v in mapping.items():
        text = text.replace(k, v)
    return text

def akcent_islandzki(text):
    text = usun_polskie_znaki(text)
    text = text.replace('ch', 'h').replace('Ch', 'H')
    text = text.replace('cz', 'ts').replace('Cz', 'Ts')
    text = text.replace('sz', 's').replace('Sz', 'S')
    text = text.replace('c', 'ts').replace('C', 'Ts')
    text = text.replace('w', 'v').replace('W', 'V')
    return sklej_pojedyncze_litery(text)

def akcent_wloski(text):
    text = usun_polskie_znaki(text)
    text = text.replace('y', 'i').replace('Y', 'I')
    text = text.replace('w', 'v').replace('W', 'V')
    return sklej_pojedyncze_litery(text)

def akcent_finski(text):
    text = usun_polskie_znaki(text)
    text = text.replace('ch', 'h').replace('Ch', 'H')
    text = text.replace('cz', 'ts').replace('Cz', 'Ts')
    text = text.replace('sz', 's').replace('Sz', 'S')
    text = text.replace('rz', 'r').replace('Rz', 'R')
    text = text.replace('c', 'ts').replace('C', 'Ts')
    text = text.replace('w', 'v').replace('W', 'V')
    text = text.replace('b', 'p').replace('B', 'P')
    text = text.replace('d', 't').replace('D', 'T')
    text = text.replace('g', 'k').replace('G', 'K')
    return sklej_pojedyncze_litery(text)

def akcent_angielski(text):
    text = usun_polskie_znaki(text)
    text = text.replace('szcz', 'shch').replace('Szcz', 'Shch').replace('SZCZ', 'SHCH')
    text = text.replace('sz', 'sh').replace('Sz', 'Sh')
    text = text.replace('cz', 'ch').replace('Cz', 'Ch')
    text = text.replace('rz', 'zh').replace('Rz', 'Zh')
    text = text.replace('ch', 'h').replace('Ch', 'H')
    text = text.replace('w', 'v').replace('W', 'V')
    return sklej_pojedyncze_litery(text)

def akcent_francuski(text):
    text = usun_polskie_znaki(text)
    text = text.replace('ch', 'sh').replace('Ch', 'Sh')
    text = text.replace('sz', 'sh').replace('Sz', 'Sh')
    text = text.replace('cz', 'tch').replace('Cz', 'Tch')
    text = text.replace('rz', 'j').replace('Rz', 'J')
    text = text.replace('h', '').replace('H', '')
    text = text.replace('r', 'gh').replace('R', 'Gh')
    text = text.replace('w', 'v').replace('W', 'V')
    return sklej_pojedyncze_litery(text)

def akcent_niemiecki(text):
    text = usun_polskie_znaki(text)
    text = text.replace('sz', 'sch').replace('Sz', 'Sch')
    text = text.replace('cz', 'tsch').replace('Cz', 'Tsch')
    text = text.replace('rz', 'rsch').replace('Rz', 'Rsch')
    text = text.replace('w', 'v').replace('W', 'V')
    text = text.replace('v', 'f').replace('V', 'F')
    return sklej_pojedyncze_litery(text)

def akcent_hiszpanski(text):
    text = usun_polskie_znaki(text)
    text = text.replace('sz', 's').replace('Sz', 'S')
    text = text.replace('cz', 'ch').replace('Cz', 'Ch')
    text = text.replace('rz', 'r').replace('Rz', 'R')
    text = text.replace('w', 'b').replace('W', 'B')
    text = text.replace('v', 'b').replace('V', 'B')
    return sklej_pojedyncze_litery(text)

def procesuj_z_ochrona_tagow(text, funkcja_akcentu):
    parts = re.split(r'(<[^>]+>)', text)
    for i in range(0, len(parts), 2): 
        parts[i] = funkcja_akcentu(parts[i])
    return "".join(parts)

def oczysc_tekst_tts(tekst, z_normalizacja=True):
    if z_normalizacja:
        tekst = normalizuj_liczby(tekst)
    tekst = re.sub(r'[\*=]+', '', tekst)
    tekst = re.sub(r'^#+\s*', '', tekst, flags=re.MULTILINE)
    tekst = re.sub(r'\([^)]*\)', '', tekst)
    tekst = re.sub(r'\b(khh|hh|pff|ahh|ehh)\b[\.\s]*', '... ', tekst, flags=re.IGNORECASE)
    tekst = re.sub(r'(?i)[,\s]*z\s*wplecionymi\s*wdechami', '', tekst)
    tekst = re.sub(r'(?i)[,\s]*z\s*wdech(em|ami)', '', tekst)
    tekst = re.sub(r'^\s*,\s*', '', tekst, flags=re.MULTILINE)
    tekst = re.sub(r'([!\?\.])\s*,\s*', r'\1 ', tekst)
    tekst = re.sub(r',\s*\.\.\.', '...', tekst)
    tekst = re.sub(r'(?:\.\s*){4,}', '... ', tekst)
    tekst = re.sub(r'([!\?\.])\s*\.\.\.\s*', r'\1 ', tekst)
    tekst = re.sub(r'^\s*\.\.\.\s*', '', tekst, flags=re.MULTILINE)
    tekst = re.sub(r'\.\.\.([^\s\.])', r'... \1', tekst)
    tekst = re.sub(r' {2,}', ' ', tekst)
    return tekst.strip()


# --- 1. WCZYTYWANIE PLIKU ---
st.write("### 1. Wczytywanie pliku źródłowego")

czy_plik_wczytany = bool(st.session_state.file_content)

file_name = st.text_input(
    "Nazwa pliku (np. rozdzial.txt, strona.html):",
    disabled=czy_plik_wczytany,
    help="Aby zmienić plik, wyczyść pamięć poniżej." if czy_plik_wczytany else ""
)

col_wczytaj, col_wyczysc = st.columns(2)

if col_wczytaj.button(
    "Wczytaj plik do pamięci",
    disabled=czy_plik_wczytany,
    help="Plik jest już wczytany. Wyczyść pamięć, aby wczytać inny." if czy_plik_wczytany else ""
):
    if not file_name.strip():
        st.error("⚠️ Podaj nazwę pliku przed wczytaniem!")
    elif os.path.exists(file_name):
        _, ext = os.path.splitext(file_name)
        st.session_state.file_ext = ext.lower()
        if ext.lower() == ".docx":
            doc = docx.Document(file_name)
            st.session_state.file_content = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(file_name, "r", encoding="utf-8") as f:
                st.session_state.file_content = f.read()
        st.session_state.oryginalna_nazwa = os.path.splitext(os.path.basename(file_name))[0]
        st.session_state.plik_katalog = os.path.dirname(file_name) or "."
        st.success(f"Wczytano plik: {file_name}!")
        st.rerun()
    else:
        st.error(f"Nie znaleziono pliku: {file_name}.")

if col_wyczysc.button(
    "🗑️ Wyczyść pamięć",
    disabled=not czy_plik_wczytany,
    help="Pamięć jest już pusta." if not czy_plik_wczytany else ""
):
    st.session_state.file_content = ""
    st.session_state.file_ext = ""
    st.success("Wyczyszczono. Możesz wczytać nową treść.")
    st.rerun()

if czy_plik_wczytany:
    st.info(f"✅ Plik wczytany ({len(st.session_state.file_content)} znaków). Możesz teraz uruchomić przetwarzanie.")

# --- 2. KONFIGURACJA PRACY ---
st.write("### 2. Konfiguracja Pracy")

reczny_kod_iso = ""

if api_dostepne:
    mode = st.radio(
        "Wybierz tryb pracy:",
        (
            "🌍 Zwykły Tłumacz (Używa AI - Kosztuje kredyty)",
            "🎬 Tryb Reżysera (Używa darmowych skryptów Pythona)"
        )
    )
else:
    st.info("💡 Tryb Tłumacza AI jest wyłączony (brak pliku golden_key.env). Aktywne są wyłącznie darmowe funkcje lokalne.")
    mode = "🎬 Tryb Reżysera (Używa darmowych skryptów Pythona)"

if "Tłumacz" in mode:
    def submit_lang():
        if st.session_state.target_lang_input.strip() != "":
            st.session_state.current_lang = st.session_state.target_lang_input
            st.session_state.ready_to_translate = True
            st.session_state.target_lang_input = ""

    st.text_input("Wpisz język docelowy (np. Fiński, Islandzki, Angielski):", key="target_lang_input", on_change=submit_lang)
    target_lang = st.session_state.current_lang if st.session_state.current_lang else ""
else:
    target_lang = st.selectbox(
        "Wybierz akcent lub tryb czyszczenia (Tylko dla czytników ekranu - HTML):",
        (
            "Żaden (Czyszczenie BEZ normalizacji liczb)",
            "Żaden (Czyszczenie Z normalizacją liczb)",
            "Islandzki (np. Guðrún / eSpeak)",
            "Angielski (np. Samantha / Mark / Zira / Hazel)",
            "Francuski (np. Thomas / Amelie / Julie)",
            "Niemiecki (np. Stefan / Markus / Katja / Hedda)",
            "Hiszpański (np. Jorge / Monica / Helena)",
            "Włoski (np. Alice / Luca / Elsa)",
            "Fiński (np. Satu / Mikko / Heidi)",
            "🔧 Naprawiacz Tagów (Tylko wstrzyknięcie kodu ISO)"
        )
    )
    if "Naprawiacz" in target_lang:
        reczny_kod_iso = st.text_input("Podaj dwuliterowy kod ISO (np. en, fr, de):", max_chars=2).lower()

# --- 3. PRZETWARZANIE ---
_brak_pliku = not czy_plik_wczytany
_brak_jezyka = ("Tłumacz" in mode and not target_lang.strip()) or ("Naprawiacz" in target_lang and not reczny_kod_iso.strip())

_przetworz_disabled = _brak_pliku or _brak_jezyka
if _brak_pliku:
    _przetworz_help = "Wczytaj najpierw plik źródłowy."
elif _brak_jezyka:
    _przetworz_help = "Wpisz język docelowy lub kod ISO przed uruchomieniem."
else:
    _przetworz_help = ""

if st.button("Uruchom Przetwarzanie", disabled=_przetworz_disabled, help=_przetworz_help):
    with st.spinner("Przetwarzanie w toku..."):
        content = st.session_state.file_content
        ext = st.session_state.file_ext
        result_text = ""
        iso_code = "pl"

        oryginalna_nazwa = st.session_state.oryginalna_nazwa
        plik_katalog = st.session_state.plik_katalog
        bezpieczny_jezyk = usun_polskie_znaki(target_lang.split()[0])
        safe_lang = re.sub(r'[^a-zA-Z0-9]', '', bezpieczny_jezyk)

        if "Tłumacz" in mode and api_dostepne:
            # TRYB AI (PŁATNY)
            base_name = f"tlumaczenie_{oryginalna_nazwa}_{safe_lang}"
            temp_filename = f"temp_{base_name}.jsonl"
            
            sys_prompt = f"""# Rola
Jesteś ekspertem w dziedzinie tłumaczeń literackich i technicznych. Twoje tłumaczenia są wzorcowe pod względem jakości, naturalności i wierności oryginałowi.

## Zadanie
Przetłumacz **cały** dostarczony tekst na język: **{target_lang}**.

## Zasady jakości (obowiązkowe)
- Tłumaczenie musi być **dokładne**, **naturalne** i **zachowywać styl oraz ton oryginału**.
- Zachowaj strukturę akapitów i podział na linie.
- Imiona własne, nazwy geograficzne i terminy specjalistyczne tłumacz zgodnie z konwencją języka docelowego.
- Jeśli oryginał zawiera humor, idiomy lub metafory — oddaj ich **sens i efekt**, nie dosłowne brzmienie.

## Zasady techniczne (krytyczne)
- **BEZWZGLĘDNIE** zachowaj wszystkie znaczniki HTML i Markdown bez żadnych modyfikacji.
- Jeśli tekst zawiera HTML, tłumacz **WYŁĄCZNIE** tekst widoczny dla użytkownika (zawartość tagów), nigdy atrybuty ani nazwy tagów.
- Nie dodawaj żadnych komentarzy, wyjaśnień ani wstępów od siebie.

## Format odpowiedzi
Zwróć **wyłącznie** przetłumaczony tekst — nic więcej."""

            akapity = content.split('\n')
            bloki = []
            obecny_blok = ""

            for akapit in akapity:
                if len(obecny_blok) + len(akapit) < 10000:
                    obecny_blok += akapit + "\n"
                else:
                    if obecny_blok.strip():
                        bloki.append(obecny_blok.strip())
                    obecny_blok = akapit + "\n"
            if obecny_blok.strip():
                bloki.append(obecny_blok.strip())
            
            if len(bloki) > 1 and len(bloki[-1]) < 4000:
                if len(bloki[-2]) + len(bloki[-1]) < 16000:
                    bloki[-2] += "\n\n" + bloki[-1]
                    bloki.pop()

            wczytane_bloki = {}
            if os.path.exists(temp_filename):
                st.info("💡 Wykryto plik zapisu. Odtwarzanie opłaconego postępu...")
                try:
                    with open(temp_filename, "r", encoding="utf-8") as f:
                        for line in f:
                            if line.strip():
                                data = json.loads(line)
                                wczytane_bloki[data['id']] = data['text']
                except Exception as e:
                    wyswietl_blad_ai(e, "⚠️ Błąd odczytu pliku tymczasowego:")

            przetlumaczono_wszystko = True
            progress_bar = st.progress(0)

            # Pętla tłumaczenia
            for i, blok in enumerate(bloki):
                if i in wczytane_bloki:
                    st.write(f"✅ Blok {i+1} odzyskany z pliku zapisu.")
                    progress_bar.progress((i + 1) / len(bloki))
                    continue

                st.write(f"⏳ Tłumaczenie bloku {i+1} z {len(bloki)}... ({len(blok)} znaków)")
                
                messages_payload = [
                    {"role": "system", "content": sys_prompt}
                ]
                
                if i > 0 and (i - 1) in wczytane_bloki:
                    messages_payload.append({
                        "role": "assistant", 
                        "content": wczytane_bloki[i-1]
                    })
                    user_content = f"[KRYTYCZNE: Kontynuuj tłumaczenie poniższego tekstu. Zachowaj absolutną spójność terminologii (w tym nazw własnych), tonu oraz ogólnego stylu (literackiego lub technicznego) z Twoją poprzednią odpowiedzią. Nie zmieniaj narzuconego wcześniej charakteru tekstu.]\n\n{blok}"
                else:
                    user_content = blok
                    
                messages_payload.append({"role": "user", "content": user_content})

                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages_payload,
                        temperature=0.3
                    )
                    
                    przetlumaczony_fragment = response.choices[0].message.content.strip()
                    wczytane_bloki[i] = przetlumaczony_fragment
                    
                    with open(temp_filename, "a", encoding="utf-8") as f:
                        f.write(json.dumps({"id": i, "text": przetlumaczony_fragment}, ensure_ascii=False) + "\n")

                    progress_bar.progress((i + 1) / len(bloki))

                except openai.RateLimitError:
                    st.error(f"🚨 BRAK ŚRODKÓW LUB LIMIT API! Przerwano na bloku {i+1}.")
                    przetlumaczono_wszystko = False
                    break
                except Exception as e:
                    wyswietl_blad_ai(e)
                    przetlumaczono_wszystko = False
                    break

            dostepne_teksty = [wczytane_bloki[idx] for idx in sorted(wczytane_bloki.keys())]
            result_text = "\n\n".join(dostepne_teksty).strip()

            if not przetlumaczono_wszystko:
                st.warning(f"⚠️ **Praca wstrzymana.** Stan został bezpiecznie zapisany w pliku `{temp_filename}`. Zasil konto API i wczytaj oryginał ponownie, by kontynuować od tego miejsca.")
                st.write("### Podgląd opłaconego dotąd fragmentu:")
                st.text_area("Częściowe tłumaczenie (tylko do podglądu, nie musisz kopiować):", value=result_text, height=350)
                st.stop()
            else:
                st.write("⏳ Ostatni krok: Generowanie architektury języka dla czytników ekranu...")
                iso_prompt = f"""Podaj **wyłącznie** dwuliterowy kod języka ISO 639-1 dla języka: **{target_lang}**. Odpowiedź musi zawierać tylko i wyłącznie dwuliterowy kod, np.: `fi`, `it`, `en`."""
                try:
                    lang_code_response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": iso_prompt}],
                        temperature=0.0
                    )
                    iso_code_raw = lang_code_response.choices[0].message.content.strip().lower()
                    iso_code = re.sub(r'[^a-z]', '', iso_code_raw)
                    
                    if not iso_code or len(iso_code) > 3:
                        raise ValueError(f"Model AI zwrócił nieobsługiwany format: {iso_code_raw}")
                        
                except Exception as e:
                    wyswietl_blad_ai(e, "⚠️ Nie udało się automatycznie pobrać kodu ISO. Użyto domyślnego tagu 'pl'. Użyj naprawiacza języka w trybie reżysera, by wstrzyknąć poprawny tag językowy.")
                    iso_code = "pl"
                
                if os.path.exists(temp_filename):
                    os.remove(temp_filename)

        else:
            # TRYB REŻYSERA (DARMOWY)
            if "Naprawiacz" in target_lang:
                base_name = f"naprawiony_{oryginalna_nazwa}_{reczny_kod_iso}"
                result_text = content
                iso_code = reczny_kod_iso if reczny_kod_iso else "pl"
            elif "Żaden" in target_lang or safe_lang.lower() == "aden":
                base_name = f"oczyszczony_{oryginalna_nazwa}"
                iso_code = "pl"
            else:
                base_name = f"akcent_{oryginalna_nazwa}_{safe_lang}"
                iso_code = "pl"
                
                if "Angielski" in target_lang:
                    iso_code = "en"
                elif "Francuski" in target_lang:
                    iso_code = "fr"
                elif "Niemiecki" in target_lang:
                    iso_code = "de"
                elif "Hiszpa" in target_lang:
                    iso_code = "es"
                else:
                    iso_code = "pl"

            if "Naprawiacz" not in target_lang:
                try:
                    detected_lang = detect(content)
                    if detected_lang != 'pl':
                        st.warning("⚠️ Uwaga: Wykryto język główny inny niż polski! Reguły fonetyczne są przystosowane do polszczyzny. Efekt może być nieprzewidywalny.")
                except LangDetectException:
                    pass

                if "Islandzki" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_islandzki)
                elif "Włoski" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_wloski)
                elif "Fiński" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_finski)
                elif "Angielski" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_angielski)
                elif "Francuski" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_francuski)
                elif "Niemiecki" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_niemiecki)
                elif "Hiszpa" in target_lang:
                    content = oczysc_tekst_tts(content, z_normalizacja=True)
                    result_text = procesuj_z_ochrona_tagow(content, akcent_hiszpanski)
                elif "Żaden" in target_lang:
                    if "BEZ" in target_lang:
                        result_text = oczysc_tekst_tts(content, z_normalizacja=False)
                    else:
                        result_text = oczysc_tekst_tts(content, z_normalizacja=True)

        # --- ZAPIS DO PLIKU (WSPÓLNY) ---
        if result_text:
            if ext == ".docx":
                out_filename = os.path.join(plik_katalog, f"{base_name}.docx")
                
                if "Reżysera" in mode:
                    doc = docx.Document(file_name)
                    for p in doc.paragraphs:
                        for run in p.runs:
                            if run.text.strip():
                                if "Islandzki" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_islandzki)
                                elif "Włoski" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_wloski)
                                elif "Fiński" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_finski)
                                elif "Angielski" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_angielski)
                                elif "Francuski" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_francuski)
                                elif "Niemiecki" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_niemiecki)
                                elif "Hiszpa" in target_lang: run.text = procesuj_z_ochrona_tagow(oczysc_tekst_tts(run.text, True), akcent_hiszpanski)
                                elif "Żaden" in target_lang and "Naprawiacz" not in target_lang: 
                                    run.text = oczysc_tekst_tts(run.text, "BEZ" not in target_lang)
                                
                                if "Naprawiacz" in target_lang:
                                    rPr = run._r.get_or_add_rPr()
                                    lang = rPr.find(qn('w:lang'))
                                    if lang is None:
                                        lang = OxmlElement('w:lang')
                                        rPr.append(lang)
                                    lang.set(qn('w:val'), iso_code)

                    doc.save(out_filename)
                    st.write("### Wynik Przetwarzania:")
                    st.text_area("Podgląd (płaski tekst):", value=result_text, height=350)
                    st.success(f"🎉 Plik .docx przetworzony i zapisany jako: {out_filename}")
                    
                else: 
                    st.write("### Wynik Przetwarzania:")
                    st.text_area("Podgląd (płaski tekst):", value=result_text, height=350)
                    nowy_doc = docx.Document()
                    for linia in result_text.split('\n'):
                        p = nowy_doc.add_paragraph(linia)
                        for run in p.runs:
                            rPr = run._r.get_or_add_rPr()
                            lang = rPr.find(qn('w:lang'))
                            if lang is None:
                                lang = OxmlElement('w:lang')
                                rPr.append(lang)
                            lang.set(qn('w:val'), iso_code)
                    nowy_doc.save(out_filename)
                    st.success(f"🎉 Plik przetłumaczony i zapisany jako: {out_filename} (docx z tagiem lang={iso_code})")

            elif ext in ['.html', '.htm']:
                out_filename = os.path.join(plik_katalog, f"{base_name}{ext}")
                if 'lang=' in result_text.lower():
                    result_text = re.sub(r'(<html[^>]*?)lang=["\'][^"\']+["\']', fr'\1lang="{iso_code}"', result_text, flags=re.IGNORECASE)
                elif '<html' in result_text.lower():
                    result_text = re.sub(r'(<html[^>]*)>', fr'\1 lang="{iso_code}">', result_text, flags=re.IGNORECASE)
                with open(out_filename, "w", encoding="utf-8") as f:
                    f.write(result_text)
                st.write("### Wynik Przetwarzania:")
                st.text_area("Skopiuj tekst poniżej:", value=result_text, height=350)
                st.success(f"🎉 Plik zapisany jako: {out_filename} (lang={iso_code})")

            elif ext in ['.txt', '.md'] and ("Tłumacz" in mode or "Naprawiacz" in target_lang):
                out_filename = os.path.join(plik_katalog, f"{base_name}.html")
                linie = result_text.split('\n')
                tytul = linie[0].strip() if linie and linie[0].strip() else "Dokument"
                przetworzony_tekst = result_text.replace('\n', '<br>\n')
                html_content = f"""<!DOCTYPE html>\n<html lang="{iso_code}">\n<head>\n<meta charset="utf-8">\n<title>{tytul}</title>\n</head>\n<body>\n{przetworzony_tekst}\n</body>\n</html>"""
                with open(out_filename, "w", encoding="utf-8") as f:
                    f.write(html_content)
                st.write("### Wynik Przetwarzania:")
                st.text_area("Skopiuj tekst poniżej:", value=result_text, height=350)
                st.success(f"🎉 Plik zapisany jako: {out_filename} (Otrzymał format HTML z tagiem lang={iso_code})")

            else:
                out_filename = os.path.join(plik_katalog, f"{base_name}{ext if ext else '.txt'}")
                with open(out_filename, "w", encoding="utf-8") as f:
                    f.write(result_text)
                st.write("### Wynik Przetwarzania:")
                st.text_area("Skopiuj tekst poniżej:", value=result_text, height=350)
                st.success(f"🎉 Plik zapisany jako: {out_filename}")
