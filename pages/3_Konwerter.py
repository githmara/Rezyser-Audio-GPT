import streamlit as st
import streamlit.components.v1 as components
import docx
import os
import re

st.set_page_config(page_title="Architekt Audiobooka", layout="wide")

st.markdown("""
    <style>
        header {visibility: hidden !important;}
        footer {visibility: hidden !important;}
        #MainMenu {visibility: hidden !important;}
    </style>
""", unsafe_allow_html=True)

components.html("""
    <script>
        window.parent.document.documentElement.lang = 'pl';
    </script>
""", width=0, height=0)

st.title("Architekt Audiobooków (Tylko Struktura)")
st.info("Ten moduł służy do przygotowania architektury pliku dla czytników ekranu (nawigacja klawiszami 1 i h po nagłówkach) oraz ElevenLabs. Zmienia słowa kluczowe (Akt, Rozdział, Prolog, Epilog) na Nagłówki pierwszego poziomu, pozostawiając całą resztę skryptu bez zmian.")

file_name = st.text_input("Nazwa pliku źródłowego (np. surowa_ksiazka.txt, gotowy_skrypt.docx):")

_brak_nazwy = not file_name.strip()
_plik_nie_istnieje = not _brak_nazwy and not os.path.exists(file_name)

_btn_disabled = _brak_nazwy or _plik_nie_istnieje
if _brak_nazwy:
    _btn_help = "Podaj nazwę pliku źródłowego."
elif _plik_nie_istnieje:
    _btn_help = f"Nie znaleziono pliku: {file_name}"
else:
    _btn_help = ""

if st.button("Buduj Architekturę dla ElevenLabs", disabled=_btn_disabled, help=_btn_help):
    # Odczyt pliku
    if file_name.endswith(".docx"):
        doc_in = docx.Document(file_name)
        tekst = "\n".join([p.text for p in doc_in.paragraphs])
    else:
        with open(file_name, "r", encoding="utf-8") as f:
            tekst = f.read()

    nowy_doc = docx.Document()
    nowy_doc.core_properties.author = "Reżyser"
    nowy_doc.core_properties.comments = ""

    for linia in tekst.splitlines():
        linia = linia.strip()
        if not linia:
            continue
        # Usuwanie tagów HTML (jeśli tekst pobrano z przeglądarki lub użyto poligloty na zwykłym tekście)
        linia = re.sub(r'<[^>]+>', '', linia)

        linia = linia.strip()
        if not linia:
            continue
        # Usuwanie znaczników nagłówków z Markdowna (np. ### lub ####)
        linia = re.sub(r'^#+\s*', '', linia)

        # Detekcja nagłówków głównych (Tną plik na rozdziały w ElevenLabs)
        if re.match(r"^[=\-\s]*(Czołówka|Rozdział|Prolog|Epilog|Akt)", linia, re.IGNORECASE):
            czysty_naglowek = re.sub(r'^[=\-\s]+|[=\-\s]+$', '', linia)
            nowy_doc.add_heading(czysty_naglowek, level=1)

        # Detekcja scen (Zwykły pogrubiony tekst, by nie robić śmietnika w spisie treści)
        elif re.match(r"^[=\-\s]*Scena", linia, re.IGNORECASE):
            czysty_naglowek = re.sub(r'^[=\-\s]+|[=\-\s]+$', '', linia)
            p = nowy_doc.add_paragraph()
            run = p.add_run(czysty_naglowek)
            run.bold = True

        else:
            # Zwykły, oczyszczony tekst dialogu lub opisu
            nowy_doc.add_paragraph(linia)

    katalog = os.path.dirname(file_name) or "."
    oryginalna_nazwa = os.path.splitext(os.path.basename(file_name))[0]
    out_name = os.path.join(katalog, f"architektura_{oryginalna_nazwa}.docx")
    nowy_doc.save(out_name)
    st.success(f"Perfetto! Plik gotowy. Nagłówki pierwszego poziomu zostały ustawione. Zapisano jako: {out_name}")
