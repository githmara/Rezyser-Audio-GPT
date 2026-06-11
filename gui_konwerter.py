"""
gui_konwerter.py – Panel modułu „Architekt Audiobooków".

Zastępuje pages/3_Konwerter.py (Streamlit).
Dziedziczy po wx.Panel; podpinany do MainFrame z main.py.

Wersja 13.1: cały tekst widoczny dla użytkownika przechodzi przez
:mod:`i18n` (klucze z ``dictionaries/pl/gui/ui.yaml`` – sekcja ``konwerter``).
"""

import os
import re
from functools import lru_cache
from pathlib import Path

import docx
import wx
import yaml

import sciezki
from i18n import t


# Co ile tur konwerter wstawia nagłówek H1 „Scena N" (cięcie rozdziału w
# ElevenLabs). 5 tur ≈ 5-7,5 tys. znaków, czyli kilka minut audio na scenę —
# zbliżone do naturalnego rozdziału audiobooka. Krótszy próg sztucznie tnie
# narrację; dłuższy traci nawigację dla NVDA.
TURY_NA_SCENE = 5

# Katalog `dictionaries/` w roocie repo (lub obok exe po zamrożeniu) — źródło
# wszystkich słów-kluczy, którymi konwerter rozpoznaje strukturę dokumentu.
_DICTIONARIES_DIR = sciezki.KATALOG_BAZOWY / "dictionaries"


def _wyluskaj_slowo_tury(fmt: str) -> str:
    """Z formatu ``"\\n\\n--- Tura {numer} ---\\n\\n"`` wyłuskuje słowo ``"Tura"``.

    Bierze fragment przed ``{numer}``, zdejmuje myślniki i białe znaki. Zwraca
    ``""`` gdy formatu nie da się sparsować (brak ``{numer}``). Dzięki temu
    słowo nagłówka tury jest pochodną istniejącego klucza
    ``opowiesci.tura_naglowek_format`` — nie trzeba go dublować osobnym kluczem
    ani hardkodować listy języków w Pythonie.
    """
    przed, sep, _ = fmt.partition("{numer}")
    if not sep:
        return ""
    return przed.strip().strip("-").strip()


@lru_cache(maxsize=1)
def _slowa_kluczowe_konwertera() -> dict[str, frozenset[str]]:
    """Skanuje ``ui.yaml`` WSZYSTKICH paczek raz i zwraca zbiory słów-kluczy.

    Tło: konwerter rozpoznaje trzy rodzaje markerów struktury dokumentu, a
    wszystkie pochodzą z tłumaczeń w ``dictionaries/<kod>/gui/ui.yaml``:

    * ``"tura"``     – słowo nagłówka tury Opowieści, wyłuskane z
      ``opowiesci.tura_naglowek_format`` (``"--- Tura {numer} ---"`` → ``Tura``);
    * ``"rozdzial"`` – słowa nagłówków H1 tnących rozdział (``konwerter.
      naglowki_rozdzialow``);
    * ``"scena"``    – słowa nagłówków scen pogrubianych bez wpisu w TOC
      (``konwerter.naglowki_scen``).

    Unia ze WSZYSTKICH paczek, BEZ detekcji języka pliku wejściowego — to
    świadoma, ugruntowana filozofia tego modułu (D1/D2 = A, v17.2.2): konwerter
    nie zna języka pliku (user mógł zmienić język aplikacji w trakcie projektu,
    plik bywa mieszany albo zbyt krótki dla detektora), więc rozpoznaje słowa
    z dowolnej paczki. Dodanie 10. języka = nowy folder + wypełnione ``ui.yaml``,
    zero edycji tego pliku (zero-touch). Jeden przebieg po YAML-ach
    (``lru_cache``) zasila wszystkie trzy regexy poniżej.
    """
    zbiory: dict[str, set[str]] = {
        "tura": set(), "rozdzial": set(), "scena": set(),
    }
    if not _DICTIONARIES_DIR.is_dir():
        return {k: frozenset(v) for k, v in zbiory.items()}
    for jezyk_dir in sorted(_DICTIONARIES_DIR.iterdir()):
        ui_path = jezyk_dir / "gui" / "ui.yaml"
        if not ui_path.is_file():
            continue
        try:
            with open(ui_path, "r", encoding="utf-8") as fh:
                dane = yaml.safe_load(fh) or {}
        except (OSError, yaml.YAMLError):
            continue
        opowiesci = dane.get("opowiesci") or {}
        fmt = opowiesci.get("tura_naglowek_format")
        if isinstance(fmt, str):
            slowo = _wyluskaj_slowo_tury(fmt)
            if slowo:
                zbiory["tura"].add(slowo)
        konwerter = dane.get("konwerter") or {}
        for klucz, cel in (("naglowki_rozdzialow", "rozdzial"),
                           ("naglowki_scen", "scena")):
            lista = konwerter.get(klucz)
            if isinstance(lista, list):
                for slowo in lista:
                    if isinstance(slowo, str) and slowo.strip():
                        zbiory[cel].add(slowo.strip())
    return {k: frozenset(v) for k, v in zbiory.items()}


def _alternatywa(slowa: frozenset[str]) -> str:
    """Buduje bezpieczną alternatywę regexa ``(?:a|b|c)`` z escapowanych słów.

    Sortowanie po długości malejąco jest kosmetyczne (dla dopasowania boolowego
    kolejność alternatyw nie ma znaczenia), ale czyni regex deterministycznym.
    """
    return "|".join(re.escape(s) for s in sorted(slowa, key=len, reverse=True))


@lru_cache(maxsize=1)
def _regex_tury() -> re.Pattern:
    """Regex znacznika tury ``"--- <słowo> N ---"`` ze słów wszystkich paczek."""
    slowa = _slowa_kluczowe_konwertera()["tura"]
    if not slowa:
        return re.compile(r"$^")
    return re.compile(
        r"^---\s*(?:" + _alternatywa(slowa) + r")\s+(\d+)\s*---$",
        re.IGNORECASE,
    )


@lru_cache(maxsize=1)
def _regex_naglowkow_rozdzialow() -> re.Pattern:
    """Regex nagłówków H1 (Rozdział/Prolog/Epilog/Akt/… ze wszystkich paczek).

    Dopasowanie prefiksowe po opcjonalnych ozdobnikach ``=-`` i białych znakach
    (np. ``"=== Rozdział 1 ==="`` albo ``"Rozdział 1: Początek"``).
    """
    slowa = _slowa_kluczowe_konwertera()["rozdzial"]
    if not slowa:
        return re.compile(r"$^")
    return re.compile(r"^[=\-\s]*(?:" + _alternatywa(slowa) + r")", re.IGNORECASE)


@lru_cache(maxsize=1)
def _regex_scen() -> re.Pattern:
    """Regex nagłówków scen (Scena/Scene/Szene/… ze wszystkich paczek)."""
    slowa = _slowa_kluczowe_konwertera()["scena"]
    if not slowa:
        return re.compile(r"$^")
    return re.compile(r"^[=\-\s]*(?:" + _alternatywa(slowa) + r")", re.IGNORECASE)


class KonwerterPanel(wx.Panel):
    """
    Panel narzędzia „Architekt Audiobooków".

    Funkcjonalność:
        - Przyjmuje ścieżkę do pliku .txt lub .docx
        - Przetwarza tekst: czyści HTML/Markdown, wykrywa nagłówki
          (Czołówka, Rozdział, Prolog, Epilog, Akt) i sceny
        - Zapisuje wynik jako architektura_<oryginalna_nazwa>.docx
          w tym samym katalogu co plik źródłowy
        - Sukces / błąd raportuje przez wx.MessageBox (A11y)
    """

    def __init__(self, parent: wx.Window) -> None:
        super().__init__(parent, style=wx.TAB_TRAVERSAL)
        self.SetName(t("konwerter.panel_name"))

        self._build_ui()
        self._bind_events()

    # ------------------------------------------------------------------
    # Budowanie interfejsu
    # ------------------------------------------------------------------
    def _build_ui(self) -> None:
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        # --- Nagłówek narzędzia ---
        heading = wx.StaticText(self, label=t("konwerter.heading"))
        heading_font = heading.GetFont()
        heading_font.SetPointSize(16)
        heading_font.MakeBold()
        heading.SetFont(heading_font)

        # --- Opis narzędzia ---
        description = wx.TextCtrl(
            self,
            value=t("konwerter.tool_description"),
            style=wx.TE_MULTILINE | wx.TE_READONLY | wx.NO_BORDER,
        )
        # Upodabniamy tło pola do tła głównego okna, żeby nie wyglądało jak pole do wpisywania
        description.SetBackgroundColour(self.GetBackgroundColour())

        # --- Separator ---
        separator = wx.StaticLine(self)

        # --- Etykieta + pole wejściowe na nazwę / ścieżkę pliku ---
        lbl_file = wx.StaticText(self, label=t("konwerter.lbl_plik"))

        self._txt_file = wx.TextCtrl(
            self,
            style=wx.TE_PROCESS_ENTER,
            name=t("konwerter.txt_plik_name"),
        )
        self._txt_file.SetHint(t("konwerter.txt_plik_hint"))

        self._btn_browse = wx.Button(self, label=t("konwerter.btn_przegladaj"))
        self._btn_browse.SetToolTip(t("konwerter.btn_przegladaj_tooltip"))

        # Poziomy sizer: pole tekstowe (rozszerzalne) + przycisk po prawej
        file_row_sizer = wx.BoxSizer(wx.HORIZONTAL)
        file_row_sizer.Add(
            self._txt_file,
            proportion=1,
            flag=wx.ALIGN_CENTER_VERTICAL | wx.RIGHT,
            border=6,
        )
        file_row_sizer.Add(
            self._btn_browse,
            flag=wx.ALIGN_CENTER_VERTICAL,
        )

        # --- Przycisk akcji ---
        self._btn_build = wx.Button(self, label=t("konwerter.btn_buduj"))
        self._btn_build.SetToolTip(t("konwerter.btn_buduj_tooltip"))

        # --- Złożenie layoutu ---
        main_sizer.Add(heading,       flag=wx.ALL, border=16)
        main_sizer.Add(description,   flag=wx.LEFT | wx.RIGHT | wx.BOTTOM, border=16)
        main_sizer.Add(separator,     flag=wx.EXPAND | wx.LEFT | wx.RIGHT, border=16)
        main_sizer.Add(lbl_file,      flag=wx.LEFT | wx.TOP | wx.RIGHT, border=16)
        main_sizer.Add(
            file_row_sizer,
            flag=wx.EXPAND | wx.LEFT | wx.RIGHT | wx.TOP,
            border=8,
        )
        main_sizer.Add(
            self._btn_build,
            flag=wx.LEFT | wx.TOP | wx.BOTTOM,
            border=16,
        )

        self.SetSizer(main_sizer)

    # ------------------------------------------------------------------
    # Podpięcie zdarzeń
    # ------------------------------------------------------------------
    def _bind_events(self) -> None:
        self._btn_build.Bind(wx.EVT_BUTTON, self._on_build)
        self._btn_browse.Bind(wx.EVT_BUTTON, self._on_browse)
        # Enter w polu tekstowym też uruchamia akcję
        self._txt_file.Bind(wx.EVT_TEXT_ENTER, self._on_build)

    # ------------------------------------------------------------------
    # Otwieranie okna wyboru pliku
    # ------------------------------------------------------------------
    def _on_browse(self, _event: wx.Event) -> None:
        """Otwiera systemowy dialog wyboru pliku i wstawia ścieżkę do pola."""
        with wx.FileDialog(
            self,
            message=t("konwerter.file_dlg_title"),
            wildcard=t("konwerter.file_dlg_wildcard"),
            style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST,
        ) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                self._txt_file.SetValue(dlg.GetPath())
                self._txt_file.SetFocus()

    # ------------------------------------------------------------------
    # Logika przetwarzania (przeniesiona z 3_Konwerter.py)
    # ------------------------------------------------------------------
    def _on_build(self, _event: wx.Event) -> None:
        """Obsługuje kliknięcie przycisku „Buduj Architekturę"."""
        file_name = self._txt_file.GetValue().strip()

        # --- Walidacja wejścia ---
        if not file_name:
            wx.MessageBox(
                t("konwerter.brak_pliku_tresc"),
                t("konwerter.brak_pliku_tytul"),
                wx.OK | wx.ICON_WARNING,
                self,
            )
            self._txt_file.SetFocus()
            return

        if not os.path.exists(file_name):
            wx.MessageBox(
                t("konwerter.plik_nie_istnieje_tresc", sciezka_pliku=file_name),
                t("konwerter.plik_nie_istnieje_tytul"),
                wx.OK | wx.ICON_ERROR,
                self,
            )
            self._txt_file.SetFocus()
            return

        # --- Odczyt pliku źródłowego ---
        try:
            if file_name.lower().endswith(".docx"):
                doc_in = docx.Document(file_name)
                tekst = "\n".join(p.text for p in doc_in.paragraphs)
            else:
                with open(file_name, "r", encoding="utf-8") as fh:
                    tekst = fh.read()
        except Exception as exc:
            wx.MessageBox(
                t("konwerter.blad_odczytu_tresc", tresc_bledu=str(exc)),
                t("konwerter.blad_odczytu_tytul"),
                wx.OK | wx.ICON_ERROR,
                self,
            )
            return

        # --- Przetwarzanie treści ---
        # Pole `author` .docx jest widoczne dla NVDA w dymkach podpowiedzi
        # Eksploratora – bierzemy je z i18n, żeby nie pokazywać polskiego
        # „Reżyser" użytkownikom, którzy aplikację mają na innym języku.
        nowy_doc = docx.Document()
        nowy_doc.core_properties.author = t("konwerter.author_metadata")
        nowy_doc.core_properties.comments = ""

        # v15.1: liczniki dla trybu Opowieści — co TURY_NA_SCENE tur wstawiamy
        # H1 „Scena N", pozostałe znaczniki `--- Tura N ---` strippujemy
        # (meta-info „Tura 7" w audiobooku łamie immersję). Licznik `tury_od_h1`
        # liczy się od ostatniego wstawionego H1 i resetuje przy każdym nowym H1.
        # Inicjalizacja `tury_od_h1 = TURY_NA_SCENE` żeby pierwsza tura
        # dokumentu wymusiła H1 „Scena 1" (warunek `>=` spełniony od razu).
        tury_od_h1 = TURY_NA_SCENE
        scena_counter = 0
        # Trzy regexy zasilane jednym przebiegiem po ui.yaml wszystkich paczek
        # (unia słów-kluczy, bez detekcji języka pliku — patrz
        # `_slowa_kluczowe_konwertera`). Liczone raz przed pętlą.
        regex_tura = _regex_tury()
        regex_rozdzial = _regex_naglowkow_rozdzialow()
        regex_scena = _regex_scen()

        def _wstaw_h1_sceny() -> None:
            nonlocal scena_counter, tury_od_h1
            scena_counter += 1
            tury_od_h1 = 0
            etykieta = t("konwerter.scena_naglowek_format", numer=scena_counter)
            nowy_doc.add_heading(etykieta, level=1)

        for linia in tekst.splitlines():
            linia = linia.strip()
            if not linia:
                continue

            # Usuwanie tagów HTML
            linia = re.sub(r'<[^>]+>', '', linia).strip()
            if not linia:
                continue

            # Usuwanie znaczników nagłówków Markdown (np. ### lub ####)
            linia = re.sub(r'^#+\s*', '', linia)

            # v15.1: Detekcja znaczników tury z Opowieści.
            # Pierwsza tura w dokumencie i każda po TURY_NA_SCENE turach
            # II-osobowych od ostatniego H1 dostają osobną H1 „Scena N".
            # Same znaczniki `--- Tura N ---` są strippowane.
            if regex_tura.match(linia):
                if tury_od_h1 >= TURY_NA_SCENE:
                    _wstaw_h1_sceny()
                tury_od_h1 += 1
                continue

            # Detekcja nagłówków głównych (tnących plik na rozdziały w ElevenLabs).
            # Słowa-klucze (Rozdział/Prolog/Epilog/Akt/…) z `konwerter.
            # naglowki_rozdzialow` WSZYSTKICH paczek — pełna obsługa 9 języków,
            # zero-touch dla 10. (v17.2.2).
            if regex_rozdzial.match(linia):
                czysty = re.sub(r'^[=\-\s]+|[=\-\s]+$', '', linia)
                nowy_doc.add_heading(czysty, level=1)

            # Detekcja scen (pogrubiony tekst, bez wpisu w spisie treści).
            # Słowa-klucze (Scena/Scene/Szene/…) z `konwerter.naglowki_scen`
            # WSZYSTKICH paczek (v17.2.2).
            elif regex_scena.match(linia):
                czysty = re.sub(r'^[=\-\s]+|[=\-\s]+$', '', linia)
                p = nowy_doc.add_paragraph()
                run = p.add_run(czysty)
                run.bold = True

            else:
                nowy_doc.add_paragraph(linia)

        # --- Zapis pliku wynikowego ---
        katalog = os.path.dirname(os.path.abspath(file_name))
        oryginalna_nazwa = os.path.splitext(os.path.basename(file_name))[0]
        out_name = os.path.join(katalog, f"architektura_{oryginalna_nazwa}.docx")

        try:
            nowy_doc.save(out_name)
        except Exception as exc:
            wx.MessageBox(
                t("konwerter.blad_zapisu_tresc", tresc_bledu=str(exc)),
                t("konwerter.blad_zapisu_tytul"),
                wx.OK | wx.ICON_ERROR,
                self,
            )
            return

        wx.MessageBox(
            t("konwerter.sukces_tresc", sciezka_pliku=out_name),
            t("konwerter.sukces_tytul"),
            wx.OK | wx.ICON_INFORMATION,
            self,
        )
