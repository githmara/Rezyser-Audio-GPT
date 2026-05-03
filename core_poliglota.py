"""
core_poliglota.py – Silnik modułu „Poliglota AI".

Cały Python-owy „brain" modułu Poliglota: ładowanie reguł z plików YAML
w folderze ``dictionaries/`` i ich stosowanie wg wskazanego trybu i języka.

Publiczne API (prosty, wysokopoziomowy interfejs używany przez GUI):

    import core_poliglota

    # lista dostępnych wariantów (do wypełnienia ComboBoxa w GUI):
    core_poliglota.lista_wariantow(tryb="Rezyser",  jezyk="pl")
    core_poliglota.lista_wariantow(tryb="Szyfrant", jezyk="pl")

    # przetwarzanie tekstu:
    wynik = core_poliglota.przetworz(
        tekst,
        tryb="Szyfrant",       # lub "Rezyser"
        jezyk="pl",
        wariant="cezar",       # id z YAML (np. "islandzki", "odwracanie")
        przesuniecie=7,        # parametr zależny od algorytmu
    )

    # metadane pomocne przy zapisie pliku wynikowego:
    core_poliglota.kod_iso(tryb="Rezyser", jezyk="pl",
                           wariant="islandzki", opcje={})
    core_poliglota.sufiks_nazwy_pliku(tryb, jezyk, wariant,
                                      oryginalna_nazwa, opcje)
    core_poliglota.zapisz_wynik(...)          # HTML / DOCX / TXT z tagiem lang

Tłumacz AI (OpenAI) znajduje się w osobnym module: ``tlumacz_ai.py``.

Konwencja nazewnicza w YAML-ach (dictionaries/):

    dictionaries/
    └── <jezyk>/                          # "pl" (docelowo też "en", "de", …)
        ├── podstawy.yaml                  # polskie_znaki + alfabet
        ├── akcenty/                       # Tryb Reżysera
        │   └── <id>.yaml                  # np. islandzki.yaml
        └── szyfry/                        # Tryb Szyfranta
            └── <id>.yaml                  # np. odwracanie.yaml

Silnik skanuje te foldery leniwie (cache) i nie wymaga żadnej rejestracji
nowych plików w kodzie – wystarczy wrzucić YAML i uruchomić aplikację.
"""

from __future__ import annotations

import os
import random
import re
from typing import Any, Callable

import yaml

# python-docx potrzebne tylko do zapisu .docx – importujemy globalnie,
# bo GUI i tak ma tę zależność.
import docx
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from num2words import num2words

# 13.5: detekcja języka oparta na ``lingua-language-detector``.
# Lingua jest deterministyczna z założenia (operuje na n-gramowych modelach
# statystycznych, nie na losowych próbkach), znacznie dokładniejsza dla
# krótkich tekstów niż dawne ``langdetect``, i — co dla nas kluczowe — pozwala
# zawęzić zestaw rozpoznawanych języków do tych, dla których faktycznie mamy
# słowniki w ``dictionaries/``. Dzięki temu detektor nigdy nie zwróci kodu
# języka, którego silnik i tak nie umiałby przetworzyć.
#
# Builder ładuje modele leniwie przy pierwszej detekcji (~1–2 s, ~100 MB RAM),
# dlatego trzymamy go za lazy singletonem ``_zbuduj_detektor_lingua``,
# uruchamianym dopiero przy realnym wywołaniu, nie w czasie importu modułu.
try:
    from lingua import Language as _LinguaLanguage
    from lingua import LanguageDetectorBuilder as _LinguaBuilder
except ImportError:                                             # pragma: no cover
    _LinguaLanguage = None   # type: ignore[assignment]
    _LinguaBuilder = None    # type: ignore[assignment]


# =============================================================================
# Ścieżki, stałe i cache
# =============================================================================
_ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
DICTIONARIES_DIR = os.path.join(_ROOT_DIR, "dictionaries")

# Mapowanie nazw trybów → podfolderów języka
TRYB_REZYSER = "Rezyser"
TRYB_SZYFRANT = "Szyfrant"

_FOLDER_DLA_TRYBU: dict[str, str] = {
    TRYB_REZYSER:  "akcenty",
    TRYB_SZYFRANT: "szyfry",
}

# Cache wczytanych danych (thread-safe dla odczytu – yaml.safe_load zwraca kopię)
_CACHE_PODSTAWY:  dict[str, dict]       = {}          # jezyk → dict
_CACHE_WARIANTOW: dict[tuple[str, str], list[dict]] = {}   # (tryb, jezyk) → lista


# =============================================================================
# Funkcje niskiego poziomu – czyste, bezstanowe, używane przez akcenty i szyfry
# =============================================================================

def normalizuj_liczby(tekst: str, jezyk: str = "pl") -> str:
    """Zamienia cyfrowe zapisy liczb na słowa (np. ``123`` → ``sto dwadzieścia trzy``).

    13.3: parametr ``jezyk`` decyduje, w jakim języku ``num2words`` rozwija
    cyfrę. Domyślnie ``"pl"`` — backward-compat. Akcent angielski musi
    przekazać ``"en"`` (``123`` → ``one hundred and twenty-three``), bo
    inaczej w angielski tekst wkleilibyśmy polskie słowa.

    Liczby, których biblioteka nie potrafi zapisać dla danego języka
    (nieznane locale, bardzo duże wartości), zostawiamy w oryginale —
    lepiej zostawić cyfry niż wybuchnąć.
    """
    def zamien(match: re.Match[str]) -> str:
        try:
            return num2words(match.group(), lang=jezyk)
        except Exception:
            return match.group()
    return re.sub(r"\d+", zamien, tekst)


def sklej_pojedyncze_litery(tekst: str) -> str:
    """Scala wiszące pojedyncze litery oddzielone spacją (np. „w y s” → „wys”)."""
    return re.sub(r"(?i)\b([a-z])\s+", r"\1", tekst)


def oczysc_tekst_tts(tekst: str, z_normalizacja: bool = True,
                     jezyk: str = "pl") -> str:
    """Oczyszcza tekst pod syntezator mowy (TTS).

    Usuwa:
      * bełkot onomatopeiczny („khh”, „pff”, „ahh”, …),
      * gwiazdki, znaki `=`, znaczniki Markdown (nagłówki),
      * nawiasy kwadratowe z przypisami reżyserskimi,
      * wielokrotne kropki i spacje,
      * frazy typu „z wplecionymi wdechami” (artefakty gpt-4).

    Jeśli ``z_normalizacja`` jest prawdziwe – dodatkowo zamienia cyfry
    na słowa w języku ``jezyk`` (por. :func:`normalizuj_liczby`).
    """
    if z_normalizacja:
        tekst = normalizuj_liczby(tekst, jezyk)
    tekst = re.sub(r"[\*=]+", "", tekst)
    tekst = re.sub(r"^#+\s*", "", tekst, flags=re.MULTILINE)
    tekst = re.sub(r"\([^)]*\)", "", tekst)
    tekst = re.sub(r"\b(khh|hh|pff|ahh|ehh)\b[\.\s]*", "... ", tekst, flags=re.IGNORECASE)
    tekst = re.sub(r"(?i)[,\s]*z\s*wplecionymi\s*wdechami", "", tekst)
    tekst = re.sub(r"(?i)[,\s]*z\s*wdech(em|ami)", "", tekst)
    tekst = re.sub(r"^\s*,\s*", "", tekst, flags=re.MULTILINE)
    tekst = re.sub(r"([!\?\.])\s*,\s*", r"\1 ", tekst)
    tekst = re.sub(r",\s*\.\.\.", "...", tekst)
    tekst = re.sub(r"(?:\.\s*){4,}", "... ", tekst)
    tekst = re.sub(r"([!\?\.])\s*\.\.\.\s*", r"\1 ", tekst)
    tekst = re.sub(r"^\s*\.\.\.\s*", "", tekst, flags=re.MULTILINE)
    tekst = re.sub(r"\.\.\.([^\s\.])", r"... \1", tekst)
    tekst = re.sub(r" {2,}", " ", tekst)
    return tekst.strip()


def procesuj_z_ochrona_tagow(tekst: str, funkcja: Callable[[str], str]) -> str:
    """Stosuje ``funkcja`` tylko do zwykłego tekstu, pomijając tagi HTML.

    Dzieli wejście na naprzemienne fragmenty „tekst” / „<tag>”; funkcja
    przetwarzająca trafia wyłącznie na pozycje parzyste listy.
    """
    parts = re.split(r"(<[^>]+>)", tekst)
    for i in range(0, len(parts), 2):
        parts[i] = funkcja(parts[i])
    return "".join(parts)


def _zastosuj_zamiany(tekst: str, zamiany: list[dict]) -> str:
    """Stosuje listę par ``{wzor, zamiana, regex?}`` z pliku YAML.

    Wzory oznaczone ``regex: true`` używają ``re.sub``, pozostałe są
    traktowane jako zwykłe stringi i zamieniane przez ``str.replace``.
    """
    for para in zamiany:
        wzor    = para.get("wzor", "")
        zamiana = para.get("zamiana", "")
        if para.get("regex"):
            tekst = re.sub(wzor, zamiana, tekst)
        else:
            tekst = tekst.replace(wzor, zamiana)
    return tekst


def _usun_polskie_znaki(tekst: str, podstawy: dict) -> str:
    """Transliteruje polskie diakrytyki wg listy z ``podstawy.yaml``.

    Funkcja NIE normalizuje liczb samodzielnie – to robi pipeline akcentu
    (flaga ``normalizuj_liczby`` w YAML-u akcentu).
    """
    return _zastosuj_zamiany(tekst, podstawy.get("polskie_znaki", []))


# =============================================================================
# Ładowanie plików YAML
# =============================================================================

def _zaladuj_yaml(sciezka: str) -> dict:
    """Wczytuje pojedynczy plik YAML i zwraca słownik (lub ``{}``)."""
    try:
        with open(sciezka, "r", encoding="utf-8") as fh:
            data = yaml.safe_load(fh)
    except Exception as exc:  # noqa: BLE001
        print(f"[core_poliglota] Błąd wczytywania {sciezka}: {exc}")
        return {}
    return data if isinstance(data, dict) else {}


def _zaladuj_podstawy(jezyk: str) -> dict:
    """Zwraca dict z ``<jezyk>/podstawy.yaml`` (cache w pamięci)."""
    if jezyk in _CACHE_PODSTAWY:
        return _CACHE_PODSTAWY[jezyk]

    sciezka = os.path.join(DICTIONARIES_DIR, jezyk, "podstawy.yaml")
    if not os.path.exists(sciezka):
        print(f"[core_poliglota] Brak pliku podstaw dla jezyka {jezyk}: {sciezka}")
        _CACHE_PODSTAWY[jezyk] = {}
        return _CACHE_PODSTAWY[jezyk]

    _CACHE_PODSTAWY[jezyk] = _zaladuj_yaml(sciezka)
    return _CACHE_PODSTAWY[jezyk]


def slowa_akcentu(jezyk: str) -> list[str]:
    """Zwraca listę słów-wyzwalaczy parsera akcentów dla danego języka.

    13.3+: pole ``slowo_akcent`` w ``dictionaries/<jezyk>/podstawy.yaml``
    zawiera listę słów (lower-case), które ``core_rezyser`` traktuje jako
    znacznik „tu mowa o akcencie X" w Księdze Świata. Funkcja zwraca:

      * listę z YAML-a, gdy pole istnieje i jest niepuste,
      * fallback ``["akcent"]`` dla starszych paczek bez pola lub gdy
        plik podstaw nie istnieje (zachowanie sprzed 13.3).

    Wynik jest płaską listą stringów; wpisy nie-stringowe filtrujemy
    defensywnie (gdyby ktoś wpisał liczbę albo zagnieżdżoną listę).
    """
    podstawy = _zaladuj_podstawy(jezyk)
    surowe = podstawy.get("slowo_akcent")
    if not isinstance(surowe, list):
        return ["akcent"]
    czyste = [str(s).strip().lower() for s in surowe if isinstance(s, str) and s.strip()]
    return czyste or ["akcent"]


def _zaladuj_warianty(tryb: str, jezyk: str) -> list[dict]:
    """Zwraca listę wszystkich wariantów (akcentów/szyfrów) dla pary tryb+język.

    Rezultat jest sortowany wg pola ``kolejnosc`` (rosnąco) z YAML-i,
    a jako tie-breaker używana jest etykieta (alfabetycznie).
    """
    klucz = (tryb, jezyk)
    if klucz in _CACHE_WARIANTOW:
        return _CACHE_WARIANTOW[klucz]

    podfolder = _FOLDER_DLA_TRYBU.get(tryb)
    if podfolder is None:
        _CACHE_WARIANTOW[klucz] = []
        return []

    katalog = os.path.join(DICTIONARIES_DIR, jezyk, podfolder)
    if not os.path.isdir(katalog):
        _CACHE_WARIANTOW[klucz] = []
        return []

    warianty: list[dict] = []
    for nazwa_pliku in os.listdir(katalog):
        if not nazwa_pliku.lower().endswith((".yaml", ".yml")):
            continue
        cfg = _zaladuj_yaml(os.path.join(katalog, nazwa_pliku))
        if not cfg:
            continue
        if "id" not in cfg:
            cfg["id"] = os.path.splitext(nazwa_pliku)[0]
        if "etykieta" not in cfg:
            cfg["etykieta"] = cfg["id"]
        cfg.setdefault("kolejnosc", 999)
        warianty.append(cfg)

    warianty.sort(key=lambda c: (c.get("kolejnosc", 999), c.get("etykieta", "")))
    _CACHE_WARIANTOW[klucz] = warianty
    return warianty


# =============================================================================
# Publiczne API – listowanie i wyszukiwanie wariantów
# =============================================================================

def lista_wariantow(tryb: str, jezyk: str = "pl") -> list[dict]:
    """Zwraca listę wariantów do wypełnienia ComboBox w GUI.

    Każdy element to słownik z *przynajmniej* kluczami ``id``, ``etykieta``,
    ``opis``, ``iso``, ``kategoria``. GUI zazwyczaj interesuje tylko
    ``etykieta`` (widoczna w liście) i ``id`` (przekazywane do
    :func:`przetworz`). Pozostałe pola mogą być przydatne w tooltipach.
    """
    return list(_zaladuj_warianty(tryb, jezyk))


def wariant_po_id(tryb: str, jezyk: str, id_: str) -> dict | None:
    """Zwraca surową konfigurację wariantu (z YAML) po jego ``id``, lub ``None``."""
    for cfg in _zaladuj_warianty(tryb, jezyk):
        if cfg.get("id") == id_:
            return cfg
    return None


def wariant_po_etykiecie(tryb: str, jezyk: str, etykieta: str) -> dict | None:
    """Zwraca surową konfigurację wariantu po jego ``etykieta`` (widocznej w GUI)."""
    for cfg in _zaladuj_warianty(tryb, jezyk):
        if cfg.get("etykieta") == etykieta:
            return cfg
    return None


# =============================================================================
# Publiczne API – detekcja języka tekstu źródłowego (multi-language ready)
# =============================================================================
#
# Kontekst: dziś GUI Poligloty hardkoduje ``JEZYK_BAZOWY = "pl"`` i wywołuje
# ``langdetect.detect()`` tylko do ostrzegania użytkownika. Gdy powstaną
# drugie, trzecie `dictionaries/<kod>/`, GUI będzie musiał podmienić hardkod
# na wynik :func:`wykryj_jezyk_zrodlowy` – infrastruktura jest już gotowa.

def _jezyk_kompletny(kod: str) -> bool:
    """Czy folder ``dictionaries/<kod>/`` ma komplet plików pełnej obsługi?

    Kryterium kompletności (spójne z TODO § 4 krok 1):

      1. ``podstawy.yaml``            – alfabet + transliteracja (Cezar)
      2. ``gui/ui.yaml``              – tłumaczenie warstwy interfejsu
      3. ``akcenty/<id>.yaml`` ≥ 1   – tryb Reżysera
      4. ``szyfry/<id>.yaml``  ≥ 1   – tryb Szyfranta

    Folder ``rezyser/`` z trybami AI jest świadomie POMIJANY — w 13.x
    to wciąż polskie prompty `gpt-4o`, nie kontrakt każdego języka
    (do reanalizy w 14.x, gdy tryby Reżysera dostaną wielojęzyczne
    prompty systemowe).

    Args:
        kod: dwuliterowy kod języka (nazwa folderu w ``dictionaries/``).

    Returns:
        True gdy wszystkie cztery warunki spełnione, False przy stubach.
    """
    folder = os.path.join(DICTIONARIES_DIR, kod)
    if not os.path.isfile(os.path.join(folder, "podstawy.yaml")):
        return False
    if not os.path.isfile(os.path.join(folder, "gui", "ui.yaml")):
        return False
    for pod in ("akcenty", "szyfry"):
        pod_dir = os.path.join(folder, pod)
        if not os.path.isdir(pod_dir):
            return False
        if not any(p.endswith(".yaml") for p in os.listdir(pod_dir)):
            return False
    return True


def dostepne_jezyki_bazowe() -> list[str]:
    """Zwraca posortowaną listę kodów KOMPLETNYCH języków w ``dictionaries/``.

    „Kompletny" oznacza folder spełniający wszystkie cztery warunki
    z :func:`_jezyk_kompletny`. Stuby (np. folder z samym `podstawy.yaml`
    i `gui/ui.yaml`, ale bez `akcenty/` i `szyfry/`) są filtrowane —
    silnik nie umiałby przetwarzać tekstu w takim języku, więc nie powinny
    pojawiać się w komunikatach typu „obsługiwane języki" ani w selektorze
    języka interfejsu w GUI.

    Skutek dla v13.1: dziś tylko `pl` przechodzi filtr. Każdy kolejny
    release minor 13.x (zgodnie z `TODO_wielojezycznosc.md` § 4)
    dorzuca jeden nowy folder w pełni wdrożony, więc lista rośnie o jedną
    pozycję per release — bez zmian w kodzie Pythona.

    Returns:
        np. ``["pl"]`` po 13.1, a po wdrożeniu fińskiego w 13.2 → ``["fi", "pl"]``.
    """
    if not os.path.isdir(DICTIONARIES_DIR):
        return []
    wynik: list[str] = []
    for nazwa in sorted(os.listdir(DICTIONARIES_DIR)):
        if not os.path.isdir(os.path.join(DICTIONARIES_DIR, nazwa)):
            continue
        if _jezyk_kompletny(nazwa):
            wynik.append(nazwa)
    return wynik


def natywna_nazwa(kod: str) -> str:
    """Natywna nazwa języka (prefiks `etykieta` w `<kod>/podstawy.yaml`).

    Przykład: dla ``kod="fi"`` zwraca ``"Suomi"`` (z etykiety
    ``"Suomi – foneettiset perusteet"``). Fallback na sam kod ISO,
    gdy etykieta nie ma separatora ` – ` lub nie istnieje.

    13.4: wyciągnięte z ``main._natywna_nazwa`` na poziom modułu, żeby GUI
    Poligloty mogło użyć tego helpera w komunikacie A11Y o zmianie języka
    pipeline'u (NVDA odczytuje pełne natywne nazwy zamiast kodów ISO).
    """
    etyk = _zaladuj_podstawy(kod).get("etykieta", "")
    if isinstance(etyk, str) and etyk:
        nazwa = etyk.split(" – ", 1)[0].strip()
        if nazwa:
            return nazwa
    return kod


def lista_wspieranych_jezykow_natywnie(jezyk_pierwszy: str | None = None) -> str:
    """Zwraca natywne nazwy wspieranych języków, gotowe do komunikatu GUI.

    Skanuje `dictionaries/<kod>/podstawy.yaml`, czyta pole `etykieta`
    i bierze prefiks przed separatorem ` – ` (em-dash z otaczającymi
    spacjami — konwencja przyjęta we WSZYSTKICH plikach `podstawy.yaml`).
    Format zwrotny to natywne nazwy języków rozdzielone przecinkami,
    z zachowaniem oryginalnych znaków (cyrylica, Þ, Æ itp.).

    Sortowanie hybrydowe:
      * gdy ``jezyk_pierwszy`` jest podany i obecny w wyniku
        :func:`dostepne_jezyki_bazowe` — ten kod idzie na pierwszą
        pozycję, reszta alfabetycznie po kodzie ISO. Pozwala GUI
        priorytetyzować język interfejsu użytkownika w komunikatach
        (np. dla użytkownika EN: „English, Polski, Suomi…" zamiast
        twardego „Polski, …").
      * w przeciwnym razie — PL twardo na pierwszym miejscu (rdzeń
        projektu, bezpieczny domyślny), reszta po ISO.

    Args:
        jezyk_pierwszy: Opcjonalny dwuliterowy kod ISO języka, który
                        ma się pojawić jako pierwszy element listy.
                        Najczęściej `i18n.aktualny_jezyk()`.
                        Jeśli None lub kod nieobecny w `dictionaries/` —
                        spadamy na PL-hardcode.

    Returns:
        Po 13.1: ``"Polski"`` (jedyny w pełni wdrożony język).
        Po 13.2 z fińskim, gdy `jezyk_pierwszy="fi"`: ``"Suomi, Polski"``.
        Pusty string, gdy `dictionaries/` nie istnieje lub żaden język
        nie przechodzi filtra kompletności z :func:`dostepne_jezyki_bazowe`.
    """
    kody = dostepne_jezyki_bazowe()
    if not kody:
        return ""

    if jezyk_pierwszy and jezyk_pierwszy in kody:
        kolejnosc = [jezyk_pierwszy] + sorted(k for k in kody if k != jezyk_pierwszy)
    elif "pl" in kody:
        kolejnosc = ["pl"] + sorted(k for k in kody if k != "pl")
    else:
        kolejnosc = sorted(kody)

    natywne: list[str] = []
    for kod in kolejnosc:
        nazwa = natywna_nazwa(kod)
        if nazwa and nazwa != kod:
            natywne.append(nazwa)
    return ", ".join(natywne)


# Minimalna długość tekstu (po strip), przy której uznajemy detekcję za
# wiarygodną. Lingua dla krótszych próbek miewa fałszywe alarmy (np. „OK"
# bywa klasyfikowane jako fiński). Niżej operujemy na sticky-fallbacku: zbyt
# krótki akapit dziedziczy język po poprzednim (a na samym początku pliku —
# po parametrze ``jezyk`` przekazanym do :func:`przetworz`).
_MIN_TEKST_DLA_DETEKCJI = 20


# ---------------------------------------------------------------------------
# Mapowania ISO ↔ ``lingua.Language`` i lazy singleton detektora
# ---------------------------------------------------------------------------
# 13.4.3: mapowanie nie jest już zhardkodowane w Pythonie. Każdy
# ``dictionaries/<kod>/podstawy.yaml`` deklaruje własne pole ``lingua:``
# (wartość = nazwa enum-a ``lingua.Language``, np. ``POLISH``, ``GERMAN``).
# Dzięki temu dodanie nowego języka bazowego (de/es/fr → 13.5+) sprowadza się
# do utworzenia folderu z plikami YAML — bez zmian w kodzie Pythona, spójnie
# z obietnicą „nowy język = nowy folder", którą trzymamy też dla
# ``odswiez_rezysera`` i ``dostepne_jezyki_bazowe``.

_LINGUA_MAPOWANIE_CACHE: dict[str, Any] | None = None


def _zbuduj_mapowanie_lingua() -> dict[str, Any]:
    """Skanuje ``dictionaries/<kod>/podstawy.yaml`` i zwraca mapę ISO → ``Language``.

    Pomija języki, których ``podstawy.yaml`` nie deklaruje pola ``lingua``,
    deklaruje je pustym stringiem albo wartością nieznaną dla aktualnej
    wersji ``lingua-language-detector`` (np. literówka, nowsze enum-y, jeszcze
    nieobsługiwany przez paczkę). Wynik jest cache'owany — pierwszy skan
    woła się przy budowie detektora, kolejne wywołania są O(1).

    Wynik ``{}`` (np. brak ``lingua-py`` w środowisku albo żaden
    ``podstawy.yaml`` nie ma pola ``lingua``) prowadzi do całkowitego wyłączenia
    detektora — :func:`_wykryj_jezyk_fragmentu` zwróci wtedy każdorazowo
    ``fallback`` (czyli język aktywny w GUI).
    """
    global _LINGUA_MAPOWANIE_CACHE
    if _LINGUA_MAPOWANIE_CACHE is not None:
        return _LINGUA_MAPOWANIE_CACHE
    if _LinguaLanguage is None:
        _LINGUA_MAPOWANIE_CACHE = {}
        return _LINGUA_MAPOWANIE_CACHE

    mapa: dict[str, Any] = {}
    for kod in _jezyki_obecne_w_dictionaries():
        wartosc = _zaladuj_podstawy(kod).get("lingua")
        if not isinstance(wartosc, str) or not wartosc.strip():
            continue
        nazwa_enuma = wartosc.strip().upper()
        # `getattr` zamiast `Language[...]`, żeby nieznana nazwa nie wyrzucała
        # KeyError na każdym imporcie modułu — defensywnie pomijamy.
        kandydat = getattr(_LinguaLanguage, nazwa_enuma, None)
        if kandydat is None:
            print(f"[core_poliglota] Pole lingua: '{wartosc}' w "
                  f"dictionaries/{kod}/podstawy.yaml nie jest znaną nazwą "
                  f"`lingua.Language` — język pomijany w detektorze.")
            continue
        mapa[kod] = kandydat

    _LINGUA_MAPOWANIE_CACHE = mapa
    return _LINGUA_MAPOWANIE_CACHE


def _zbuduj_mapowanie_lingua_to_iso() -> dict[str, str]:
    """Odwrócone mapowanie ``Language.name → ISO 639-1`` na bazie tej samej mapy."""
    return {lang.name: iso for iso, lang in _zbuduj_mapowanie_lingua().items()}


def _jezyki_obecne_w_dictionaries() -> list[str]:
    """Zwraca listę kodów języków, które mają chociaż ``podstawy.yaml``.

    Rozluźniona wersja :func:`dostepne_jezyki_bazowe` — nie wymaga ``szyfry/``
    ani ``akcenty/``, a jedynie obecności pliku ``podstawy.yaml``. Używana
    tylko do nakarmienia ``LanguageDetectorBuilder`` zestawem języków, który
    chcemy rozpoznawać; rzeczywista weryfikacja kompletności reguł odbywa
    się dopiero w dyspozytorach (``_przetworz_rezyser`` / ``_przetworz_szyfrant``)
    przy pomocy :class:`BrakRegulyDlaJezykaError`.
    """
    if not os.path.isdir(DICTIONARIES_DIR):
        return []
    wynik: list[str] = []
    for nazwa in sorted(os.listdir(DICTIONARIES_DIR)):
        folder = os.path.join(DICTIONARIES_DIR, nazwa)
        if not os.path.isdir(folder):
            continue
        if os.path.isfile(os.path.join(folder, "podstawy.yaml")):
            wynik.append(nazwa)
    return wynik


_LINGUA_DETEKTOR: Any = None       # cache singletona; budowa leniwa
_LINGUA_DETEKTOR_BLD_FAILED = False  # flaga, by nie powtarzać próby budowy


def _zbuduj_detektor_lingua() -> Any:
    """Lazy singleton ``LanguageDetector`` z lingua.

    Buduje detektor przy pierwszym wywołaniu, używając zestawu języków
    z :func:`_zbuduj_mapowanie_lingua` (czyli tych ``dictionaries/<kod>/``,
    których ``podstawy.yaml`` deklaruje pole ``lingua: <NAZWA_ENUMA>``).

    Lingua wymaga ≥ 2 języków w builderze – gdy w ``dictionaries/`` jest tylko
    jeden (lub zero) język z poprawnym polem ``lingua``, zwracamy ``None``
    i wywołujący spada na ``fallback``. Zwroty ``None`` są cache'owane przez
    flagę ``_LINGUA_DETEKTOR_BLD_FAILED``, żeby nie powtarzać prób budowy
    przy każdej detekcji fragmentu.
    """
    global _LINGUA_DETEKTOR, _LINGUA_DETEKTOR_BLD_FAILED

    if _LINGUA_DETEKTOR is not None:
        return _LINGUA_DETEKTOR
    if _LINGUA_DETEKTOR_BLD_FAILED:
        return None
    if _LinguaBuilder is None:
        _LINGUA_DETEKTOR_BLD_FAILED = True
        return None

    mapowanie = _zbuduj_mapowanie_lingua()
    if len(mapowanie) < 2:
        _LINGUA_DETEKTOR_BLD_FAILED = True
        return None

    _LINGUA_DETEKTOR = _LinguaBuilder.from_languages(*mapowanie.values()).build()
    return _LINGUA_DETEKTOR


def _wykryj_jezyk_fragmentu(tekst: str, fallback: str) -> str:
    """Wykrywa kod ISO języka pojedynczego fragmentu (akapitu).

    Zasady (każda w innym warunku):
      1. Tekst pusty / krótszy niż ``_MIN_TEKST_DLA_DETEKCJI`` → ``fallback``
         (dla A11y: pojedyncze „OK." nie powinno przerywać przetwarzania).
      2. Brak instancji lingua (import się nie udał, < 2 języki w
         ``dictionaries/``) → ``fallback``.
      3. ``detect_language_of`` zwraca ``None`` (lingua nie jest pewna) →
         ``fallback``.
      4. Wynik mapowany przez :func:`_zbuduj_mapowanie_lingua_to_iso`. Gdy
         ``Language.name`` nie ma odpowiednika w ``dictionaries/`` (teoretycznie
         nie powinno się zdarzyć — ten sam zestaw karmił builder) → ``fallback``.
    """
    if not isinstance(tekst, str) or len(tekst.strip()) < _MIN_TEKST_DLA_DETEKCJI:
        return fallback

    detektor = _zbuduj_detektor_lingua()
    if detektor is None:
        return fallback

    wynik = detektor.detect_language_of(tekst)
    if wynik is None:
        return fallback

    iso = _zbuduj_mapowanie_lingua_to_iso().get(wynik.name)
    if not iso:
        return fallback
    return iso if iso in _jezyki_obecne_w_dictionaries() else fallback


# ---------------------------------------------------------------------------
# Wyjątek: brak reguły dla wykrytego języka fragmentu
# ---------------------------------------------------------------------------

class BrakRegulyDlaJezykaError(RuntimeError):
    """Lingua wykryła w tekście język, dla którego brakuje żądanej reguły.

    Podnoszony przez :func:`_przetworz_rezyser` / :func:`_przetworz_szyfrant`,
    gdy w tekście wejściowym znajdzie się akapit w języku ``L``, a w
    ``dictionaries/L/<podfolder>/<wariant>.yaml`` nie istnieje plik z regułą.
    GUI łapie ten wyjątek osobno i wyświetla **długi techniczny komunikat
    w ``wx.Dialog`` z ``TextCtrl`` ``TE_READONLY``** (zgodnie z konwencją
    A11y: krótkie powiadomienia → ``wx.MessageBox``, długie techniczne →
    ``wx.Dialog`` z polem do skopiowania).

    Atrybuty:
        jezyk_kod:        Kod ISO 639-1 wykrytego języka (np. ``"ru"``).
        jezyk_natywna:    Nazwa języka w jego natywnym brzmieniu
                          (np. ``"Русский"``); jeżeli nie ma ``podstawy.yaml``
                          — równe ``jezyk_kod``.
        tryb:             ``"Rezyser"`` lub ``"Szyfrant"``.
        wariant:          ``id`` wariantu (np. ``"jakanie"``).
        oczekiwany_folder: Względna ścieżka brakującej reguły, np.
                          ``"dictionaries/ru/szyfry"``. Pomaga użytkownikowi
                          natychmiast trafić do miejsca, w którym powinien
                          dorzucić plik YAML.
    """

    def __init__(self, jezyk_kod: str, jezyk_natywna: str, tryb: str,
                 wariant: str, oczekiwany_folder: str) -> None:
        self.jezyk_kod = jezyk_kod
        self.jezyk_natywna = jezyk_natywna
        self.tryb = tryb
        self.wariant = wariant
        self.oczekiwany_folder = oczekiwany_folder
        komunikat = (
            f"Wykryto fragment w języku {jezyk_natywna} (kod '{jezyk_kod}'), "
            f"ale brakuje reguły '{wariant}' w folderze {oczekiwany_folder}.\n\n"
            f"Utwórz plik {oczekiwany_folder}/{wariant}.yaml przed "
            f"kontynuowaniem przetwarzania, albo usuń fragmenty w tym "
            f"języku z tekstu wejściowego."
        )
        super().__init__(komunikat)


# ---------------------------------------------------------------------------
# Segmentacja tekstu na akapity z ochroną tagów HTML
# ---------------------------------------------------------------------------

# Segment to trójka ``(jezyk_iso, tresc, czy_przetwarzac)``:
#   * ``jezyk_iso``       – kod języka (sticky-fallback dla krótkich akapitów),
#   * ``tresc``           – dosłowny fragment tekstu (akapit, separator, tag),
#   * ``czy_przetwarzac`` – ``True`` tylko dla właściwych akapitów tekstowych;
#                            ``False`` dla tagów HTML i separatorów ``\n\n``,
#                            które należy przepisać 1:1 do wyniku.
Segment = tuple[str, str, bool]


def _segmentuj_z_ochrona_tagow(tekst: str, fallback_jezyk: str) -> list[Segment]:
    """Dzieli tekst na segmenty z wykrytym językiem; chroni tagi HTML i separatory.

    Algorytm dwuwarstwowy:
      1. ``re.split(r"(<[^>]+>)", tekst)`` → naprzemienne pozycje tekst/tag.
         Tagi (indeksy nieparzyste) trafiają do wyniku z flagą ``False``
         i nie są nigdy poddawane detekcji ani transformacjom.
      2. Każdy fragment-tekst (indeksy parzyste) dzielimy ponownie:
         ``re.split(r"(\\n\\s*\\n)", ...)`` → akapity i pomiędzy nimi
         dosłowne separatory ``\\n\\s*\\n``. Separatory zachowujemy 1:1.
      3. Dla każdego niepustego akapitu wywołujemy
         :func:`_wykryj_jezyk_fragmentu` z fallbackiem na język poprzedniego
         akapitu – tym samym krótka linia („Tak.", „OK.") dziedziczy język
         po sąsiadach zamiast wymuszać reset na ``fallback_jezyk``.

    Pierwszy akapit (gdy nie ma jeszcze „poprzedniego") używa
    ``fallback_jezyk`` – zwykle parametr ``jezyk`` przekazany do
    :func:`przetworz`, czyli język aktywny w GUI.

    Zwraca listę gotową do iteracji w dyspozytorach – wszystko jest tam
    gotowe: kod języka per akapit, dosłowna treść, flaga „przetwarzaj".
    """
    if not isinstance(tekst, str) or not tekst:
        return []

    czesci = re.split(r"(<[^>]+>)", tekst)
    wynik: list[Segment] = []
    poprzedni_jezyk = fallback_jezyk

    for i, czesc in enumerate(czesci):
        if i % 2 == 1:
            # Tag HTML – zachowaj 1:1, nie analizuj.
            wynik.append((poprzedni_jezyk, czesc, False))
            continue
        if not czesc:
            continue
        # Drugi poziom: akapity (\n\n) z zachowaniem separatora.
        akapity = re.split(r"(\n\s*\n)", czesc)
        for j, akapit in enumerate(akapity):
            if j % 2 == 1:
                # Separator między akapitami – nie tłumaczymy.
                wynik.append((poprzedni_jezyk, akapit, False))
                continue
            if not akapit:
                continue
            jez = _wykryj_jezyk_fragmentu(akapit, fallback=poprzedni_jezyk)
            wynik.append((jez, akapit, True))
            poprzedni_jezyk = jez

    return wynik


def wykryj_jezyk_zrodlowy(
    tekst: str,
    *,
    fallback: str = "pl",
    dostepne: list[str] | None = None,
) -> str:
    """Wykrywa kod języka tekstu; waliduje wynik wobec ``dictionaries/``.

    Funkcja jest „konserwatywna" – zwraca ``fallback`` w każdym z wypadków,
    w których wynik detekcji byłby niemiarodajny:

      1. ``lingua`` nie zostało zainstalowane (brak importu na starcie),
      2. ``tekst`` jest za krótki (<``_MIN_TEKST_DLA_DETEKCJI`` znaków
         po strip),
      3. detektor zwraca ``None`` (zbyt mało sygnału do klasyfikacji),
      4. wykryty kod NIE występuje w liście ``dostepne`` – nawet jeśli
         lingua trafiła, GUI musi pozostać przy języku, który ma komplet
         działających słowników (``szyfry/`` + ``akcenty/`` + ``gui/``).

    Dzięki punktowi (4) funkcja jest „multi-language ready": dziś zawsze
    zwraca ``"pl"`` (bo to jedyny *kompletny* język), ale gdy
    ``dictionaries/en/`` dostanie szyfry, zacznie zwracać ``"en"`` dla
    angielskich tekstów.

    Args:
        tekst:    Tekst do zbadania (zwykle wczytana zawartość pliku).
        fallback: Co zwrócić, gdy detekcja się nie powiedzie lub wynik nie
                  ma swojego folderu w ``dictionaries/`` (domyślnie ``"pl"``).
        dostepne: Opcjonalna lista dozwolonych kodów. Jeśli ``None`` – funkcja
                  sama zawoła :func:`dostepne_jezyki_bazowe`. Zdefiniowanie
                  pozwala GUI odfiltrować języki, które akurat są wyłączone.

    Returns:
        Dwuliterowy kod ISO 639-1 (np. ``"pl"``, ``"en"``, ``"de"``).

    Example:
        >>> wykryj_jezyk_zrodlowy("Ala ma kota, a kot ma Alę i wychodzą razem.")
        'pl'
        >>> wykryj_jezyk_zrodlowy("???")          # za krótki + brak liter
        'pl'
        >>> wykryj_jezyk_zrodlowy("The quick brown fox jumps over the lazy dog.",
        ...                       dostepne=["pl"])  # zawężone do PL
        'pl'
    """
    # Szybka ścieżka: za krótki tekst → fallback (bez budowania detektora).
    if not isinstance(tekst, str) or len(tekst.strip()) < _MIN_TEKST_DLA_DETEKCJI:
        return fallback

    detektor = _zbuduj_detektor_lingua()
    if detektor is None:
        return fallback

    wynik = detektor.detect_language_of(tekst)
    if wynik is None:
        return fallback

    kod_wykryty = _zbuduj_mapowanie_lingua_to_iso().get(wynik.name)
    if not kod_wykryty:
        return fallback

    if dostepne is None:
        dostepne = dostepne_jezyki_bazowe()

    return kod_wykryty if kod_wykryty in dostepne else fallback


# =============================================================================
# Tryb Reżysera – pipeline akcentu fonetycznego z pliku YAML
# =============================================================================

def _aplikuj_akcent_z_yaml(tekst: str, cfg: dict, podstawy: dict,
                           jezyk: str = "pl") -> str:
    """Uruchamia pięcioetapowy pipeline akcentu wg flag w ``cfg``.

    Etapy (wykonywane w stałej kolejności):
        1. ``czysc_tekst_tts``
        2. ``normalizuj_liczby``      (gdy nie użyto pełnego czyszczenia)
        3. ``usun_polskie_znaki``
        4. ``zamiany`` (właściwe reguły fonetyczne akcentu)
        5. ``skleja_pojedyncze_litery``

    13.3: ``jezyk`` decyduje o locale ``num2words`` — domyślnie ``"pl"``
    dla backward-compat, ale wywołujący (``_przetworz_rezyser``) przekazuje
    rzeczywisty język tekstu źródłowego.
    """
    if cfg.get("czysc_tekst_tts"):
        tekst = oczysc_tekst_tts(
            tekst,
            z_normalizacja=cfg.get("normalizuj_liczby", False),
            jezyk=jezyk,
        )
    elif cfg.get("normalizuj_liczby"):
        tekst = normalizuj_liczby(tekst, jezyk)

    if cfg.get("usun_polskie_znaki"):
        tekst = _usun_polskie_znaki(tekst, podstawy)

    tekst = _zastosuj_zamiany(tekst, cfg.get("zamiany", []))

    if cfg.get("skleja_pojedyncze_litery"):
        tekst = sklej_pojedyncze_litery(tekst)

    return tekst


# ---------------------------------------------------------------------------
# Aliasy publiczne: akcent_* (kompatybilność wsteczna z gui_rezyser.py)
# ---------------------------------------------------------------------------
# Funkcje ``akcent_<język>`` stosują WYŁĄCZNIE reguły fonetyczne danego
# akcentu (normalizacja liczb + transliteracja + zamiany + scalanie
# pojedynczych liter), bez pełnego oczyszczania TTS. Dzięki temu moduł
# Reżysera może wywoływać je punktowo na pojedynczych kwestiach dialogowych
# bez ryzyka usunięcia ich zawartości (np. gwiazdek z didaskaliów).

def zastosuj_reguly_fonetyczne(tekst: str, wariant: str,
                               jezyk: str = "pl") -> str:
    """Stosuje reguły fonetyczne wybranego akcentu – bez czyszczenia TTS.

    Równoważne staremu ``akcent_*`` z pre-refaktorowej wersji: zachowuje
    gwiazdki, hashtagi i nawiasy kwadratowe (didaskalia), zmieniając
    wyłącznie fonetykę.
    """
    cfg = wariant_po_id(TRYB_REZYSER, jezyk, wariant) or {}
    podstawy = _zaladuj_podstawy(jezyk)
    tekst = normalizuj_liczby(tekst, jezyk)
    tekst = _usun_polskie_znaki(tekst, podstawy)
    tekst = _zastosuj_zamiany(tekst, cfg.get("zamiany", []))
    return sklej_pojedyncze_litery(tekst)


# <GENEROWANE_AKCENTY_REZYSERA_START>
# UWAGA: Blok poniżej jest generowany automatycznie przez skrypt
# ``odswiez_rezysera.py``. NIE edytuj go ręcznie — edycje zostaną
# nadpisane przy najbliższym uruchomieniu skryptu (po dodaniu
# nowego pliku YAML w dictionaries/<język>/akcenty/).


def akcent_islandzki(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``islandzki`` (źródła: ``dictionaries/en/akcenty/islandzki.yaml``, ``dictionaries/fi/akcenty/islandzki.yaml``, ``dictionaries/pl/akcenty/islandzki.yaml``, ``dictionaries/ru/akcenty/islandzki.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "islandzki", jezyk)


def akcent_polski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``polski`` (źródła: ``dictionaries/en/akcenty/polski.yaml``, ``dictionaries/fi/akcenty/polski.yaml``, ``dictionaries/is/akcenty/polski.yaml``, ``dictionaries/ru/akcenty/polski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "polski", jezyk)


def akcent_angielski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``angielski`` (źródła: ``dictionaries/fi/akcenty/angielski.yaml``, ``dictionaries/is/akcenty/angielski.yaml``, ``dictionaries/pl/akcenty/angielski.yaml``, ``dictionaries/ru/akcenty/angielski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "angielski", jezyk)


def akcent_rosyjski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``rosyjski`` (źródła: ``dictionaries/en/akcenty/rosyjski.yaml``, ``dictionaries/fi/akcenty/rosyjski.yaml``, ``dictionaries/is/akcenty/rosyjski.yaml``, ``dictionaries/pl/akcenty/rosyjski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "rosyjski", jezyk)


def akcent_francuski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``francuski`` (źródła: ``dictionaries/en/akcenty/francuski.yaml``, ``dictionaries/fi/akcenty/francuski.yaml``, ``dictionaries/is/akcenty/francuski.yaml``, ``dictionaries/pl/akcenty/francuski.yaml``, ``dictionaries/ru/akcenty/francuski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "francuski", jezyk)


def akcent_niemiecki(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``niemiecki`` (źródła: ``dictionaries/en/akcenty/niemiecki.yaml``, ``dictionaries/fi/akcenty/niemiecki.yaml``, ``dictionaries/is/akcenty/niemiecki.yaml``, ``dictionaries/pl/akcenty/niemiecki.yaml``, ``dictionaries/ru/akcenty/niemiecki.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "niemiecki", jezyk)


def akcent_hiszpanski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``hiszpanski`` (źródła: ``dictionaries/en/akcenty/hiszpanski.yaml``, ``dictionaries/fi/akcenty/hiszpanski.yaml``, ``dictionaries/is/akcenty/hiszpanski.yaml``, ``dictionaries/pl/akcenty/hiszpanski.yaml``, ``dictionaries/ru/akcenty/hiszpanski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "hiszpanski", jezyk)


def akcent_wloski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``wloski`` (źródła: ``dictionaries/en/akcenty/wloski.yaml``, ``dictionaries/fi/akcenty/wloski.yaml``, ``dictionaries/is/akcenty/wloski.yaml``, ``dictionaries/pl/akcenty/wloski.yaml``, ``dictionaries/ru/akcenty/wloski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "wloski", jezyk)


def akcent_finski(tekst: str, jezyk: str = "pl") -> str:
    """Alias: reguły fonetyczne akcentu ``finski`` (źródła: ``dictionaries/en/akcenty/finski.yaml``, ``dictionaries/is/akcenty/finski.yaml``, ``dictionaries/pl/akcenty/finski.yaml``, ``dictionaries/ru/akcenty/finski.yaml``)."""
    return zastosuj_reguly_fonetyczne(tekst, "finski", jezyk)

# <GENEROWANE_AKCENTY_REZYSERA_END>


def _przetworz_rezyser(tekst: str, jezyk: str, cfg: dict, opcje: dict) -> str:
    """Rezyser: akcent / oczyszczenie / naprawiacz – z dynamiczną detekcją języka.

    13.5: silnik segmentuje wejście na akapity (z ochroną tagów HTML),
    wykrywa język każdego osobno i pobiera dla niego *własną* konfigurację
    wariantu. Brak reguły dla wykrytego języka → :class:`BrakRegulyDlaJezykaError`.

    Side-channel: do ``opcje["_segmenty_wynikowe"]`` zapisywana jest lista
    krotek ``(jezyk_iso, fragment_wynikowy, czy_przetwarzany)`` zachowująca
    kolejność z wejścia. :func:`zapisz_wynik` używa jej do wstrzyknięcia
    tagu ``lang`` per akapit bez konieczności ponownej detekcji.
    """
    kategoria = cfg.get("kategoria", "")
    wariant_id = cfg.get("id", "")

    # Naprawiacz tagów nie modyfikuje treści – tylko wstrzykiwanie ISO
    # w :func:`zapisz_wynik`. Tam też dzieje się detekcja per akapit.
    if kategoria == "naprawiacz":
        return tekst

    # Oczyszczenie: pipeline TTS niezależny od reguł YAML konkretnego języka.
    # Detekcja per akapit potrzebna tylko dla locale ``num2words``.
    if kategoria == "oczyszczenie":
        segmenty_in = _segmentuj_z_ochrona_tagow(tekst, fallback_jezyk=jezyk)
        wyniki: list[str] = []
        zapisane: list[Segment] = []
        for jez_seg, fragment, czy_przetwarzac in segmenty_in:
            if not czy_przetwarzac:
                wyniki.append(fragment)
                zapisane.append((jez_seg, fragment, False))
                continue
            wynik_fr = oczysc_tekst_tts(
                fragment,
                z_normalizacja=cfg.get("normalizuj_liczby", True),
                jezyk=jez_seg,
            )
            wyniki.append(wynik_fr)
            zapisane.append((jez_seg, wynik_fr, True))
        opcje["_segmenty_wynikowe"] = zapisane
        return "".join(wyniki)

    # Zwykły akcent – per fragment szukamy YAML-a dla wykrytego języka.
    segmenty_in = _segmentuj_z_ochrona_tagow(tekst, fallback_jezyk=jezyk)
    wyniki = []
    zapisane = []
    for jez_seg, fragment, czy_przetwarzac in segmenty_in:
        if not czy_przetwarzac:
            wyniki.append(fragment)
            zapisane.append((jez_seg, fragment, False))
            continue

        cfg_jez = wariant_po_id(TRYB_REZYSER, jez_seg, wariant_id)
        if cfg_jez is None:
            raise BrakRegulyDlaJezykaError(
                jezyk_kod=jez_seg,
                jezyk_natywna=natywna_nazwa(jez_seg),
                tryb=TRYB_REZYSER,
                wariant=wariant_id,
                oczekiwany_folder=f"dictionaries/{jez_seg}/akcenty",
            )
        podstawy_jez = _zaladuj_podstawy(jez_seg)
        wynik_fr = _aplikuj_akcent_z_yaml(fragment, cfg_jez, podstawy_jez, jez_seg)
        wyniki.append(wynik_fr)
        zapisane.append((jez_seg, wynik_fr, True))

    opcje["_segmenty_wynikowe"] = zapisane
    return "".join(wyniki)


# =============================================================================
# Tryb Szyfranta – algorytmy (parametryzowane YAML-em)
# =============================================================================
# Każdy algorytm dostaje (tekst, cfg, podstawy, opcje) i zwraca string.
# `cfg`      – słownik wczytany z <szyfr>.yaml,
# `podstawy` – słownik wczytany z <język>/podstawy.yaml,
# `opcje`    – kwargs przekazane przez GUI do :func:`przetworz` (np. przesuniecie).

# Regex łapiący pojedyncze słowa we WSZYSTKICH alfabetach Unicode (łacińskim,
# kirylickim, greckim, …). Klasa `[^\W\d_]+` to "litery dowolnego skryptu" — bez
# cyfr i bez `_`. Granice słów są implicit (klasa nie obejmuje spacji ani
# interpunkcji). 13.5: zmiana z `[a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ]+` była konieczna,
# żeby `_algo_typoglikemia` i `_algo_jakanie` zadziałały na rosyjskim
# (cyrylica). Łatka analogiczna do tej w `core_rezyser.py` z 13.3.
_REGEX_SLOWA = r"[^\W\d_]+"


def _algo_odwracanie(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Rozwija skrótowce z YAML-a, a potem odwraca każde zdanie wspak.

    Kolejność:
      1. aplikuje listę ``rozwiniecia`` (regex, case-insensitive),
      2. usuwa powtórzenia słów (np. „bardzo bardzo” → „bardzo”),
      3. dzieli tekst na zdania i każde odwraca znak po znaku,
      4. zachowuje znak interpunkcyjny na końcu i kapitalizację pierwszej
         litery nowego (odwróconego) zdania.
    """
    for para in cfg.get("rozwiniecia", []):
        tekst = re.sub(para.get("wzor", ""), para.get("zamiana", ""),
                       tekst, flags=re.IGNORECASE)

    tekst = re.sub(r"\b(\w{2,})\s+\1\b", r"\1", tekst, flags=re.IGNORECASE)
    tekst = re.sub(r" +", " ", tekst)

    def odwracaj_zdanie(zdanie: str) -> str:
        if not zdanie.strip():
            return zdanie
        znak = ""
        if zdanie[-1] in ".?!":
            znak = zdanie[-1]
            zdanie = zdanie[:-1]
        odwrocone = zdanie[::-1].lower()
        if odwrocone:
            odwrocone = odwrocone[0].upper() + odwrocone[1:]
        return odwrocone + znak

    fragmenty = re.split(r"(?<=[.!?])(\s+)", tekst)
    wynik: list[str] = []
    for frag in fragmenty:
        if frag.isspace() or not frag:
            wynik.append(frag)
        else:
            wynik.append(odwracaj_zdanie(frag))
    return "".join(wynik)


def _algo_typoglikemia(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Miesza środek każdego słowa; pierwsza i ostatnia litera pozostaje.

    Słowa krótsze niż ``cfg['min_dlugosc_slowa']`` nie są zmieniane.
    """
    min_len = int(cfg.get("min_dlugosc_slowa", 4))

    def wymieszaj(match: re.Match[str]) -> str:
        slowo = match.group(0)
        if len(slowo) < min_len:
            return slowo
        srodek = list(slowo[1:-1])
        random.shuffle(srodek)
        return slowo[0] + "".join(srodek) + slowo[-1]

    return re.sub(_REGEX_SLOWA, wymieszaj, tekst)


def _algo_samogloskowiec(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Zamienia samogłoski na ``o``, zachowując polskie zmiękczenia."""
    # Krok 1 – zmiękczenia PRZED samogłoską (regex z lookahead)
    for para in cfg.get("zmiekszenia_przed_samogloska", []):
        tekst = re.sub(para.get("wzor", ""), para.get("zamiana", ""), tekst)

    # Krok 2 – zmiękczenia PRZED spółgłoską (plain string replace)
    for para in cfg.get("zmiekszenia_przed_spolgloska", []):
        tekst = tekst.replace(para.get("wzor", ""), para.get("zamiana", ""))

    # Krok 3 – samogłoski → 'o' / 'O'
    male  = cfg.get("samogloski_male", "aeiyuąęó")
    duze  = cfg.get("samogloski_wielkie", "AEIYUĄĘÓ")
    tekst = re.sub(f"[{male}]", cfg.get("zamiana_samogloski_male", "o"),    tekst)
    tekst = re.sub(f"[{duze}]", cfg.get("zamiana_samogloski_wielkie", "O"), tekst)
    return tekst


def _algo_jakanie(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Dokleja losowe zająknięcia przed każdym dłuższym słowem.

    Parametry z YAML-a:
      ``min_dlugosc_slowa`` – krótsze słowa są pomijane,
      ``min_powtorzen`` / ``max_powtorzen`` – liczba „k-k-k” dla słowa,
      ``samogloski`` – jeśli drugi znak to samogłoska → jąkamy jedną literę,
                       w przeciwnym razie → dwie (np. „pr-pr-prysznic”).
    """
    min_len  = int(cfg.get("min_dlugosc_slowa", 3))
    min_pow  = int(cfg.get("min_powtorzen", 1))
    max_pow  = int(cfg.get("max_powtorzen", 3))
    samogl   = cfg.get("samogloski", "aeiouyąęóAEIOUYĄĘÓ")

    def zacinaj(match: re.Match[str]) -> str:
        slowo = match.group(0)
        if len(slowo) < min_len:
            return slowo
        ile_powtorzen = random.randint(min_pow, max_pow)
        prefiks = slowo[:2] if len(slowo) > min_len and slowo[1] not in samogl else slowo[0]
        powtorzenia = "-".join([prefiks.lower()] * ile_powtorzen)
        if slowo[0].isupper():
            powtorzenia = powtorzenia.capitalize()
            reszta_slowa = slowo[len(prefiks):]
            return f"{powtorzenia}-{prefiks.lower()}{reszta_slowa}"
        return f"{powtorzenia}-{slowo}"

    return re.sub(_REGEX_SLOWA, zacinaj, tekst)


def _algo_waz(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Wydłuża ``s``/``z``/``sz`` w losowy długi syk (efekt „Snecko”)."""
    min_syk = int(cfg.get("min_syk", 4))
    max_syk = int(cfg.get("max_syk", 8))
    wzor    = cfg.get("wzor_syku", "(?i)(sz|s|z)")

    def sycz(match: re.Match[str]) -> str:
        znak = match.group(0)
        ile  = random.randint(min_syk, max_syk)
        if znak.lower() == "sz":
            syk = "s" * ile + "z"
            return syk.capitalize() if znak[0].isupper() else syk
        syk = znak[0].lower() * ile
        return syk.capitalize() if znak.isupper() else syk

    return re.sub(wzor, sycz, tekst)


def _algo_cezar(tekst: str, cfg: dict, podstawy: dict, opcje: dict) -> str:
    """Klasyczny szyfr Cezara na alfabecie z ``cfg['alfabet']`` lub ``podstawy``.

    Wartość przesunięcia pobierana jest w kolejności:
      1. ``opcje['przesuniecie_faktyczne']`` (użyteczne, gdy GUI wylosowało
         przesunięcie i chce odtworzyć ten sam wynik),
      2. ``opcje['przesuniecie']`` (pole z SpinCtrl),
      3. domyślnie 0.

    Jeśli wynikowe przesunięcie wynosi 0 – losowane jest z zakresu
    ``1..len(alfabet)-1``.
    """
    alfabet = cfg.get("alfabet") or podstawy.get("alfabet") or \
              "AĄBCĆDEĘFGHIJKLŁMNŃOÓPQRSTUVWXYZŹŻ"
    n = len(alfabet)

    przes = int(opcje.get("przesuniecie_faktyczne",
                          opcje.get("przesuniecie", 0)))
    if przes == 0:
        przes = random.randint(1, n - 1)
    # Zapisz faktyczne przesunięcie z powrotem do opcji – GUI użyje tego
    # do zbudowania nazwy pliku wynikowego (patrz :func:`sufiks_nazwy_pliku`).
    opcje["przesuniecie_faktyczne"] = przes

    def przesun_znak(char: str) -> str:
        upper = char.upper()
        idx = alfabet.find(upper)
        if idx == -1:
            return char
        nowy = alfabet[(idx + przes) % n]
        return nowy if char.isupper() else nowy.lower()

    return "".join(przesun_znak(c) for c in tekst)


# =============================================================================
# Dispatcher szyfrów
# =============================================================================
_ALGORYTMY_SZYFROW: dict[str, Callable[[str, dict, dict, dict], str]] = {
    "odwracanie":     _algo_odwracanie,
    "typoglikemia":   _algo_typoglikemia,
    "samogloskowiec": _algo_samogloskowiec,
    "jakanie":        _algo_jakanie,
    "waz":            _algo_waz,
    "cezar":          _algo_cezar,
}


def _przetworz_szyfrant(tekst: str, jezyk: str, cfg: dict, opcje: dict) -> str:
    """Szyfrant: dispatcher na algorytm – z dynamiczną detekcją języka per akapit.

    13.5: dla każdego akapitu pobierane są reguły z ``dictionaries/<jezyk>/
    szyfry/<wariant>.yaml`` (gdzie ``<jezyk>`` = wynik detekcji lingua,
    ``<wariant>`` = id wybrane przez użytkownika w GUI). Brak pliku dla
    wykrytego języka → :class:`BrakRegulyDlaJezykaError`.

    Cezar: pole ``opcje["przesuniecie_faktyczne"]`` (zapisywane przez
    :func:`_algo_cezar` po pierwszym losowaniu) jest współdzielone między
    fragmentami przez referencję ``opcje`` – dzięki temu pierwszy akapit
    losuje, a kolejne reużywają tego samego przesunięcia. Każdy język ma
    własny alfabet (mod ``len(alfabet)`` różny), więc finalny shift na
    rosyjskim akapicie przy ``n=33`` da inny wynik niż na polskim akapicie
    przy ``n=32`` – to spójne z założeniem multi-language szyfrowania.
    """
    wariant_id = cfg.get("id", "")
    nazwa_algo_glob = cfg.get("algorytm", "")
    if nazwa_algo_glob and nazwa_algo_glob not in _ALGORYTMY_SZYFROW:
        raise ValueError(
            f"Nieznany algorytm szyfru: „{nazwa_algo_glob}”. "
            f"Dostępne: {sorted(_ALGORYTMY_SZYFROW)}"
        )

    segmenty_in = _segmentuj_z_ochrona_tagow(tekst, fallback_jezyk=jezyk)
    wyniki: list[str] = []
    zapisane: list[Segment] = []

    for jez_seg, fragment, czy_przetwarzac in segmenty_in:
        if not czy_przetwarzac:
            wyniki.append(fragment)
            zapisane.append((jez_seg, fragment, False))
            continue

        cfg_jez = wariant_po_id(TRYB_SZYFRANT, jez_seg, wariant_id)
        if cfg_jez is None:
            raise BrakRegulyDlaJezykaError(
                jezyk_kod=jez_seg,
                jezyk_natywna=natywna_nazwa(jez_seg),
                tryb=TRYB_SZYFRANT,
                wariant=wariant_id,
                oczekiwany_folder=f"dictionaries/{jez_seg}/szyfry",
            )
        podstawy_jez = _zaladuj_podstawy(jez_seg)

        # 13.3: normalizacja diakrytyki PRZED czyszczeniem i algorytmem.
        # 13.5: per fragment, używając podstaw języka FRAGMENTU – inaczej
        # rosyjski akapit dostałby polskie mapowania znaków.
        fragment_norm = _usun_polskie_znaki(fragment, podstawy_jez)
        fragment_czysty = oczysc_tekst_tts(fragment_norm, z_normalizacja=True,
                                            jezyk=jez_seg)

        nazwa_algo = cfg_jez.get("algorytm", "")
        funkcja = _ALGORYTMY_SZYFROW.get(nazwa_algo)
        if funkcja is None:
            raise ValueError(
                f"Nieznany algorytm szyfru: „{nazwa_algo}” "
                f"(plik dictionaries/{jez_seg}/szyfry/{wariant_id}.yaml). "
                f"Dostępne: {sorted(_ALGORYTMY_SZYFROW)}"
            )
        wynik_fr = funkcja(fragment_czysty, cfg_jez, podstawy_jez, opcje)
        wyniki.append(wynik_fr)
        zapisane.append((jez_seg, wynik_fr, True))

    opcje["_segmenty_wynikowe"] = zapisane
    return "".join(wyniki)


# =============================================================================
# Publiczne API – punkt wejścia dla GUI
# =============================================================================

def przetworz(
    tekst: str,
    tryb: str,
    jezyk: str = "pl",
    wariant: str | None = None,
    **opcje: Any,
) -> str:
    """Uruchamia wybrane przetwarzanie i zwraca gotowy tekst.

    Args:
        tekst:   Tekst źródłowy (dowolnej długości).
        tryb:    ``"Rezyser"`` lub ``"Szyfrant"``.
        jezyk:   Kod ISO 639-1 języka bazowego (domyślnie ``"pl"``).
        wariant: ``id`` z YAML-a (np. ``"islandzki"``, ``"odwracanie"``),
                 ewentualnie etykieta widoczna w GUI.
        **opcje: Dodatkowe parametry zależne od algorytmu, np.:

            ``przesuniecie`` – int, dla szyfru Cezara (0 = losuj).

    Returns:
        Przetworzony tekst jako string.

    Raises:
        ValueError: gdy tryb jest nieznany lub wariantu nie odnaleziono.
    """
    if not wariant:
        raise ValueError("Parametr `wariant` jest wymagany.")

    cfg = wariant_po_id(tryb, jezyk, wariant)
    if cfg is None:
        cfg = wariant_po_etykiecie(tryb, jezyk, wariant)
    if cfg is None:
        raise ValueError(
            f"Nie znaleziono wariantu „{wariant}” dla trybu „{tryb}” "
            f"i języka „{jezyk}”."
        )

    if tryb == TRYB_REZYSER:
        return _przetworz_rezyser(tekst, jezyk, cfg, opcje)
    if tryb == TRYB_SZYFRANT:
        return _przetworz_szyfrant(tekst, jezyk, cfg, opcje)

    raise ValueError(f"Nieznany tryb: „{tryb}”. Oczekiwano „Rezyser” lub „Szyfrant”.")


# =============================================================================
# Pomocnicze – kod ISO i nazwa pliku wynikowego
# =============================================================================

def kod_iso(tryb: str, jezyk: str, wariant: str, opcje: dict | None = None) -> str:
    """Zwraca dwuliterowy kod ISO języka dla pliku wynikowego.

    Dla naprawiacza tagów ``iso`` jest podawane ręcznie przez użytkownika
    (w ``opcje["iso_reczne"]``). Dla pozostałych wariantów pochodzi z YAML-a.
    """
    cfg = wariant_po_id(tryb, jezyk, wariant) or wariant_po_etykiecie(tryb, jezyk, wariant)
    if cfg is None:
        return jezyk

    if cfg.get("kategoria") == "naprawiacz":
        return (opcje or {}).get("iso_reczne", jezyk) or jezyk

    iso = cfg.get("iso") or jezyk
    return str(iso).strip() or jezyk


def sufiks_nazwy_pliku(
    tryb: str,
    jezyk: str,
    wariant: str,
    oryginalna_nazwa: str,
    opcje: dict | None = None,
) -> str:
    """Buduje bazową nazwę pliku wynikowego (bez rozszerzenia).

    Odtwarza nazewnictwo ze starego GUI:
      * ``oczyszczony_<oryginał>``                  – oczyszczanie
      * ``<oryginał>_akcent_<akcent>``              – akcent fonetyczny
      * ``naprawiony_<oryginał>_<iso>``             – naprawiacz tagów
      * ``<oryginał>_szyfr_<szyfr>[<±przesunięcie>]`` – szyfry

    Nie ma kropki w przyrostku – GUI doda ją razem z rozszerzeniem.
    """
    opcje = opcje or {}
    cfg   = wariant_po_id(tryb, jezyk, wariant) or wariant_po_etykiecie(tryb, jezyk, wariant)
    if cfg is None:
        return oryginalna_nazwa

    kategoria = cfg.get("kategoria", "")
    wariant_id = cfg.get("id", wariant)

    if kategoria == "naprawiacz":
        iso = (opcje.get("iso_reczne") or jezyk).strip()
        return f"naprawiony_{oryginalna_nazwa}_{iso}"

    if kategoria == "oczyszczenie":
        return f"oczyszczony_{oryginalna_nazwa}"

    if kategoria == "akcent":
        return f"{oryginalna_nazwa}_akcent_{wariant_id}"

    if kategoria == "szyfr":
        base = f"{oryginalna_nazwa}_szyfr_{wariant_id}"
        # Cezar dopisuje informację o przesunięciu (np. +7, -12)
        if cfg.get("algorytm") == "cezar":
            przes = int(opcje.get("przesuniecie_faktyczne", opcje.get("przesuniecie", 0)))
            if przes != 0:
                base = f"{base}{przes:+d}"
        return base

    return f"{oryginalna_nazwa}_{wariant_id}"


# =============================================================================
# Zapis pliku wynikowego (HTML / DOCX / TXT z tagiem lang)
# =============================================================================

# Lista tagów blokowych, dla których ma sens lokalny atrybut ``lang``.
# Wybrane spośród elementów HTML5, które typowo zawierają samodzielną
# jednostkę tekstu (czytniki ekranu przełączają silnik mowy na granicy
# zmiany ``lang`` właśnie tutaj). Inline'y (``span``, ``a``) celowo poza
# listą — generowałyby setki krótkich detekcji o niskiej wiarygodności.
_PARA_TAGS_LANG = (
    "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "li", "blockquote", "dt", "dd",
    "td", "th", "caption", "figcaption", "summary",
)


def _wstrzyknij_lang_w_pelnym_html(html_text: str, iso_fallback: str) -> str:
    """Parsuje pełnoprawny HTML i ustawia atrybut ``lang`` per element blokowy.

    13.4.3: zastępuje wcześniejszy regex (działający tylko na ``<html>``).
    BeautifulSoup buduje DOM; dla każdego elementu z :data:`_PARA_TAGS_LANG`
    wykrywamy język jego ``get_text()``-u i ustawiamy ``lang="..."`` lokalnie.
    Atrybut ``lang`` na samym ``<html>`` pozostaje globalnym fallbackiem dla
    pustych/krótkich elementów (sticky-fallback z
    :func:`_wykryj_jezyk_fragmentu`).

    Parser: ``lxml`` (preferowany — szybki i tolerancyjny dla niedomkniętego
    HTML, a w środowisku 13.4.3 gwarantowany w ``requirements.txt``). Fallback
    na wbudowany ``html.parser`` w razie braku ``lxml`` w środowisku
    deweloperskim.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:                                     # pragma: no cover
        # bs4 niedostępne — wracamy do prostego ustawienia lang w <html>.
        if "lang=" in html_text.lower():
            return re.sub(
                r'(<html[^>]*?)lang=["\'][^"\']+["\']',
                fr'\1lang="{iso_fallback}"',
                html_text, flags=re.IGNORECASE,
            )
        return re.sub(
            r"(<html[^>]*)>", fr'\1 lang="{iso_fallback}">',
            html_text, flags=re.IGNORECASE,
        )

    try:
        soup = BeautifulSoup(html_text, "lxml")
    except Exception:                                       # pragma: no cover
        soup = BeautifulSoup(html_text, "html.parser")

    if soup.html is not None:
        soup.html["lang"] = iso_fallback

    for tag_name in _PARA_TAGS_LANG:
        for el in soup.find_all(tag_name):
            tekst = el.get_text(separator=" ", strip=True)
            if not tekst:
                continue
            el["lang"] = _wykryj_jezyk_fragmentu(tekst, fallback=iso_fallback)

    return str(soup)


def _ustaw_lang_runa(run: Any, iso: str) -> None:
    """Wstrzykuje ``<w:lang w:val=iso>`` do biegu Word, tworząc ``rPr`` jeśli brak."""
    rPr = run._r.get_or_add_rPr()
    lang_el = rPr.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        rPr.append(lang_el)
    lang_el.set(qn("w:val"), iso)


def _iso_per_linia(tresc: str, segmenty: list[Segment] | None,
                   iso_fallback: str) -> list[str]:
    """Mapuje każdą linię ``tresc.split('\\n')`` na kod ISO języka.

    Strategia:
      1. Jeśli ``segmenty`` są dane (z side-channel ``opcje['_segmenty_wynikowe']``)
         i ich konkatenacja zgadza się z ``tresc`` – używamy ich. To najmocniejsze
         źródło, bo zawiera detekcję wykonaną PRZED transformacją (np. cezarem,
         który zaszyfrowałby alfabet i zafałszował detekcję na wyniku).
      2. W przeciwnym razie segmentujemy ``tresc`` na żywo (przypadek naprawiacza
         tagów oraz każdego trybu, w którym wywołujący nie podał side-channelu).
      3. Buduemy mapę offset→iso ze sticky-fallbackiem: separatory ``\\n\\s*\\n``
         dziedziczą iso po ostatnim segmencie tekstowym.
      4. Iterujemy linie, przesuwając kursor o ``len(linia)+1`` (znak ``\\n``).
    """
    if segmenty is None or "".join(s[1] for s in segmenty) != tresc:
        segmenty = _segmentuj_z_ochrona_tagow(tresc, fallback_jezyk=iso_fallback)

    # Mapa offset → iso (długość mapy == len(tresc))
    iso_per_offset: list[str] = []
    ostatni_iso = iso_fallback
    for jez, fr, czy_tekst in segmenty:
        if czy_tekst:
            ostatni_iso = jez
        iso_per_offset.extend([ostatni_iso] * len(fr))

    wynik: list[str] = []
    cursor = 0
    for linia in tresc.split("\n"):
        if 0 <= cursor < len(iso_per_offset):
            wynik.append(iso_per_offset[cursor])
        else:
            wynik.append(ostatni_iso)
        cursor += len(linia) + 1  # +1 za znak '\n'
    return wynik


def zapisz_wynik(
    tresc_wynikowa: str,
    katalog_wyjscia: str,
    base_name: str,
    ext: str,
    iso_code: str,
    tryb: str,
    wariant_cfg: dict | None,
    oryginalny_content: str,
    sciezka_oryginalu: str | None = None,
    *,
    segmenty_wynikowe: list[Segment] | None = None,
) -> str:
    """Zapisuje wynik do pliku i zwraca jego ścieżkę.

    13.5: tag ``lang`` jest wstrzykiwany **per akapit / per paragraf**, a nie
    globalnie. Dla trybów Reżysera i Szyfranta używana jest mapa
    ``segmenty_wynikowe`` (side-channel z :func:`przetworz`) – zawiera ona
    detekcję wykonaną PRZED transformacją tekstu, więc działa nawet po
    szyfrowaniu cezara czy odwracaniu zdań. Dla naprawiacza tagów detekcja
    odbywa się na bieżąco, paragraf po paragrafie, na ORYGINALNEJ treści.

    Obsługiwane rozszerzenia:
      * ``.docx`` – dokument Word z tagiem ``<w:lang w:val=iso>`` per paragraf,
      * ``.html`` / ``.htm`` – HTML; jeżeli wejście ma ``<html>``, atrybut
        ``lang`` ustawiany na ``iso_code`` (jak dotąd); w pozostałych
        przypadkach budujemy nowy dokument z ``<p lang="...">`` per akapit
        (rozdzielany ``\\n\\s*\\n``),
      * ``.txt`` / ``.md`` – konwertowane do HTML; każdy akapit dostaje
        własny ``<p lang="...">``,
      * każde inne – surowy zapis tekstu z oryginalnym rozszerzeniem.

    Args:
        tresc_wynikowa:       Przetworzony tekst do zapisu.
        katalog_wyjscia:      Katalog docelowy (zwykle: katalog pliku źródłowego).
        base_name:            Nazwa pliku bez rozszerzenia
                              (wynik :func:`sufiks_nazwy_pliku`).
        ext:                  Rozszerzenie źródła (np. ``".docx"``) – decyduje
                              o formacie wyjścia.
        iso_code:             Domyślny kod języka (fallback dla pustych /
                              krótkich akapitów oraz dla atrybutu ``<html lang>``).
        tryb:                 ``"Rezyser"`` / ``"Szyfrant"`` / ``"Tlumacz"``.
        wariant_cfg:          Konfiguracja wariantu (z YAML) – potrzebna,
                              by rozpoznać „naprawiacz tagów".
        oryginalny_content:   Treść źródłowa (wykorzystywana przez naprawiacza
                              tagów, gdy ``sciezka_oryginalu`` nie istnieje).
        sciezka_oryginalu:    Pełna ścieżka oryginału (tylko dla naprawiacza
                              ``.docx`` – kopiujemy oryginalny dokument,
                              tylko wstrzykując tag ``w:lang``).
        segmenty_wynikowe:    *Keyword-only.* Side-channel z
                              :func:`przetworz` (``opcje['_segmenty_wynikowe']``).
                              Lista krotek ``(iso, fragment, czy_tekst)`` w
                              kolejności wynikowej. ``None`` → detekcja na
                              bieżąco po treści wynikowej (przypadek naprawiacza
                              i wywołań spoza Reżysera/Szyfranta).

    Returns:
        Pełna ścieżka zapisanego pliku.
    """
    jest_naprawiacz = bool(wariant_cfg and wariant_cfg.get("kategoria") == "naprawiacz")

    # -------- DOCX ---------------------------------------------------------
    if ext == ".docx":
        out_path = os.path.join(katalog_wyjscia, f"{base_name}.docx")

        if jest_naprawiacz and sciezka_oryginalu and os.path.exists(sciezka_oryginalu):
            # Otwieramy oryginał i wstrzykujemy tag lang dynamicznie per paragraf.
            doc = docx.Document(sciezka_oryginalu)
            for para in doc.paragraphs:
                tekst_para = para.text
                jez_para = (
                    _wykryj_jezyk_fragmentu(tekst_para, fallback=iso_code)
                    if tekst_para.strip() else iso_code
                )
                for run in para.runs:
                    _ustaw_lang_runa(run, jez_para)
        else:
            # Nowy dokument: side-channel (tryb Rezyser/Szyfrant) lub live-detect
            # (naprawiacz bez pliku źródła, Tlumacz, inne wywołania).
            doc = docx.Document()
            zawartosc = oryginalny_content if jest_naprawiacz else tresc_wynikowa

            sgm = None if jest_naprawiacz else segmenty_wynikowe
            iso_lista = _iso_per_linia(zawartosc, sgm, iso_code)

            linie = zawartosc.split("\n")
            for linia, iso_lin in zip(linie, iso_lista):
                p = doc.add_paragraph(linia)
                for run in p.runs:
                    _ustaw_lang_runa(run, iso_lin)
        doc.save(out_path)
        return out_path

    # -------- HTML / HTM ---------------------------------------------------
    if ext in (".html", ".htm"):
        out_path = os.path.join(katalog_wyjscia, f"{base_name}{ext}")
        tekst = tresc_wynikowa
        ma_html = "<html" in tekst.lower()

        if ma_html:
            # 13.4.3: pełnoprawny HTML — bs4 + lxml wstrzykują ``lang``
            # per element blokowy (paragraf, nagłówek, lista, komórka),
            # zachowując resztę DOM-u. Globalny ``<html lang>`` to fallback.
            tekst = _wstrzyknij_lang_w_pelnym_html(tekst, iso_code)
        else:
            # Fragment HTML / czysty tekst — owijamy akapity (``\n\s*\n``) w
            # ``<p lang="...">`` z dynamicznym językiem.
            tekst = _zbuduj_html_z_akapitow(
                tekst,
                segmenty_wynikowe if not jest_naprawiacz else None,
                iso_code,
                z_doctype=False,
            )

        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(tekst)
        return out_path

    # -------- TXT / MD → HTML (z tagiem lang per akapit) ------------------
    if ext in (".txt", ".md"):
        out_path = os.path.join(katalog_wyjscia, f"{base_name}.html")
        linie = tresc_wynikowa.split("\n")
        tytul = linie[0].strip() if linie and linie[0].strip() else "Dokument"
        body = _zbuduj_html_z_akapitow(
            tresc_wynikowa,
            segmenty_wynikowe if not jest_naprawiacz else None,
            iso_code,
            z_doctype=False,
        )
        html = (
            f'<!DOCTYPE html>\n<html lang="{iso_code}">\n'
            f'<head>\n<meta charset="utf-8">\n<title>{tytul}</title>\n</head>\n'
            f"<body>\n{body}\n</body>\n</html>"
        )
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(html)
        return out_path

    # -------- Inne – zapis surowy -----------------------------------------
    out_path = os.path.join(katalog_wyjscia, f"{base_name}{ext if ext else '.txt'}")
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write(tresc_wynikowa)
    return out_path


def _zbuduj_html_z_akapitow(tresc: str,
                            segmenty: list[Segment] | None,
                            iso_fallback: str,
                            z_doctype: bool = False) -> str:
    """Buduje HTML, owijając każdy akapit w ``<p lang="...">``.

    Akapity są oddzielane wzorcem ``\\n\\s*\\n``. Wewnątrz akapitu pojedyncze
    ``\\n`` zostaje konwertowane na ``<br>`` (tak jak we wcześniejszym
    zachowaniu ``zapisz_wynik``). HTML-special chars (``<``, ``>``, ``&``)
    NIE są ekranowane — wynik często zawiera już własne tagi z naprawiacza
    lub akcentu, a wejście do trybów Poligloty pochodzi z zaufanego źródła
    (lokalny plik użytkownika).

    Mapowanie iso → akapit czerpie z ``segmenty`` (side-channel) gdy są
    dostępne, w przeciwnym razie segmentuje ``tresc`` na żywo.
    """
    if segmenty is None or "".join(s[1] for s in segmenty) != tresc:
        segmenty = _segmentuj_z_ochrona_tagow(tresc, fallback_jezyk=iso_fallback)

    czesci_html: list[str] = []
    biezacy_akapit: list[str] = []
    biezacy_iso = iso_fallback

    def flush_akapit() -> None:
        if not biezacy_akapit:
            return
        body = "".join(biezacy_akapit)
        if not body.strip():
            biezacy_akapit.clear()
            return
        czesci_html.append(f'<p lang="{biezacy_iso}">{body.replace(chr(10), "<br>")}</p>')
        biezacy_akapit.clear()

    for jez, fr, czy_tekst in segmenty:
        if czy_tekst:
            biezacy_iso = jez
            biezacy_akapit.append(fr)
        else:
            # Separator \n\n lub tag HTML
            if re.fullmatch(r"\n\s*\n", fr):
                flush_akapit()
            else:
                # Tag HTML w środku akapitu — zachowujemy w treści akapitu.
                biezacy_akapit.append(fr)
    flush_akapit()
    return "\n".join(czesci_html)
